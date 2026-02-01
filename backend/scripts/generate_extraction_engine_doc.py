#!/usr/bin/env python3
"""
Generate REIMS2 Extraction Engine & Data Validation documentation as .docx.
Saves to: /home/hsthind/REIMS - Documents/REIMS2_Extraction_Engine_Guide.docx
"""
import os
from pathlib import Path

def main():
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("python-docx not installed. Run: pip install python-docx")
        return 1

    out_dir = Path("/home/hsthind/REIMS - Documents")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "REIMS2_Extraction_Engine_Guide.docx"

    doc = Document()
    doc.add_heading("REIMS2 Extraction Engine & Data Validation", 0)
    doc.add_paragraph(
        "Full technical reference: tools, LLM/ML usage, validation, and end-to-end flow."
    )

    # --- 1. Overview ---
    doc.add_heading("1. Extraction Engine Overview", level=1)
    doc.add_paragraph(
        "The REIMS2 extraction engine is a multi-engine PDF extraction system that combines "
        "rule-based extractors, optional ML models (LayoutLM, EasyOCR), and LLMs (Ollama, Anthropic, OpenAI, etc.) "
        "to achieve high-accuracy structured extraction from financial documents (balance sheets, income statements, "
        "cash flow, rent roll). Results are validated by a dedicated validation service before being committed."
    )

    # --- 2. Flow diagram (text) ---
    doc.add_heading("2. End-to-End Flow Diagram", level=1)
    doc.add_paragraph("The diagram below shows how all extraction and validation tools work together from upload to completed extraction.")
    flow = (
        "PDF Upload (MinIO)\n"
        "        |\n"
        "        v\n"
        "[Pre-Flight] Create extraction_run_id & Master JSON (multi-LLM audit trail)\n"
        "        |\n"
        "        v\n"
        "[STEP 1] PDFClassifier -> document_type: digital | scanned | mixed | table_heavy\n"
        "        |\n"
        "        v\n"
        "[STEP 2] MultiEngineExtractor + QualityValidator\n"
        "         Engines: PyMuPDF, PDFPlumber, Camelot, Tesseract OCR, EasyOCR, LayoutLM\n"
        "        |\n"
        "        v\n"
        "[STEP 3] (Optional) Multi-LLM: 2-3 providers -> Master JSON candidates\n"
        "        |\n"
        "        v\n"
        "[STEP 4] TemplateExtractor + FinancialTableParser -> line_items, headers\n"
        "        |\n"
        "        v\n"
        "[STEP 5] Evidence anchoring + Multi-LLM scoring -> AUTO_ACCEPT | ESCALATE_LLM | NEEDS_REVIEW\n"
        "        |\n"
        "        v\n"
        "[STEP 6] Data insertion -> balance_sheet_data, income_statement_data, cash_flow_data, rent_roll_data\n"
        "        |\n"
        "        v\n"
        "[STEP 7] ValidationService -> ValidationRule runs -> ValidationResult, ValidationRun\n"
        "        |\n"
        "        v\n"
        "[STEP 8] Financial metrics, anomaly detection, extraction log, status = completed"
    )
    doc.add_paragraph(flow)

    # --- 2.1 STEP 3 in detail: Multi-LLM -> Master JSON candidates ---
    doc.add_heading("2.1 STEP 3 in Detail: Multi-LLM (2–3 providers → Master JSON candidates)", level=2)
    doc.add_paragraph(
        "When MULTI_LLM_EXTRACTION_ENABLED is True, the orchestrator calls 2–3 LLM providers in parallel "
        "to produce candidate structured extractions. Each candidate is stored in the Master JSON for audit and scoring."
    )
    doc.add_paragraph("When it runs:", style="Heading 3")
    doc.add_paragraph(
        "STEP 3 runs after text has been extracted (STEP 2). The extracted text (or a truncation, e.g. first 8000 characters) "
        "is sent to the multi-LLM provider along with the document type (e.g. balance_sheet, income_statement). "
        "It runs only if the feature flag is enabled and API keys are configured for at least one provider."
    )
    doc.add_paragraph("Provider selection and parallel calls:", style="Heading 3")
    doc.add_paragraph(
        "The multi-LLM provider (multi_llm_provider.generate_candidates) determines which providers are available by checking "
        "API keys (ANTHROPIC_API_KEY, OPENAI_API_KEY, PERPLEXITY_API_KEY, MISTRAL_API_KEY, GOOGLE_API_KEY). "
        "It then selects up to max_providers (default 3) from the default order: anthropic, openai, perplexity, mistral, gemini. "
        "For each selected provider, it calls the corresponding adapter (e.g. AnthropicAdapter, OpenAIAdapter) asynchronously; "
        "all calls run in parallel via asyncio.gather. Each adapter calls the provider API with a system prompt and a JSON schema, "
        "and returns an LLMCandidateResult containing: provider name, model, parsed_json (structured output), raw_response preview, "
        "latency_ms, tokens_estimate, and error (if the call failed)."
    )
    doc.add_paragraph("System prompt and schema:", style="Heading 3")
    doc.add_paragraph(
        "A document-type-specific system prompt instructs the LLM to act as an expert financial analyst, extract structured data, "
        "and return only valid JSON with no markdown or explanation. Amounts are normalized to numbers (no currency symbols; "
        "parentheses as negative). The schema enforces the expected shape: for balance_sheet, line_items (array of account_name, amount, category), "
        "total_assets, total_liabilities, total_equity; for income_statement, line_items with account_name, amount, section, plus total_revenue, "
        "total_expenses, net_operating_income; for rent_roll, units (unit_id, tenant_name, monthly_rent); for cash_flow, line_items. "
        "Adapters retry up to max_retries (e.g. 2) on malformed JSON."
    )
    doc.add_paragraph("Storing candidates in Master JSON:", style="Heading 3")
    doc.add_paragraph(
        "MasterJSONService.append_candidates_and_telemetry receives the list of LLMCandidateResult and optionally a template candidate. "
        "For each result it appends a CandidateEntry to the Master JSON candidates section: source (provider name), model, parsed_json, "
        "raw_response_preview (first 500 chars), latency_ms, tokens_estimate, error. Telemetry is updated per provider: call_count, "
        "cumulative latency_ms, tokens_estimate, error_count. The updated Master JSON (including total_latency_ms) is persisted to "
        "the ExtractionRun row (extraction_runs.master_json). Thus the Master JSON holds a full chain-of-custody of all LLM candidates "
        "for that extraction run."
    )
    doc.add_paragraph("Why STEP 3 is optional:", style="Heading 3")
    doc.add_paragraph(
        "STEP 3 is marked optional because: (1) It is gated by a config flag MULTI_LLM_EXTRACTION_ENABLED (default True in ai.py). "
        "If set to False, the orchestrator skips Multi-LLM entirely. (2) The pipeline does not depend on Multi-LLM for inserting data: "
        "structured data is produced by STEP 4 (TemplateExtractor + FinancialTableParser) from the extracted text from STEP 2. "
        "Multi-LLM only adds candidate JSONs to the Master JSON for audit, evidence anchoring, and scoring; it does not feed the DB insert. "
        "(3) If no provider has an API key, generate_candidates returns an empty list and no candidates are stored. (4) If Multi-LLM throws "
        "(e.g. timeout, API error), the orchestrator catches the exception, logs a warning, and continues with STEP 4. So extraction can complete "
        "successfully without STEP 3; STEP 3 adds extra assurance and chain-of-custody when enabled and when API keys are present."
    )
    doc.add_paragraph("Code references:", style="Heading 3")
    doc.add_paragraph("multi_llm_provider: generate_candidates_async, generate_candidates, _call_one, _default_schema_for_document_type, default_system_prompt_for_document_type.", style="List Bullet")
    doc.add_paragraph("master_json_service: append_candidates_and_telemetry.", style="List Bullet")
    doc.add_paragraph("llm_adapters: get_adapter, get_available_providers; base.LLMCandidateResult, BaseLLMAdapter.generate_structured_json.", style="List Bullet")

    # --- 2.2 STEP 5 in detail: Evidence anchoring + Multi-LLM scoring ---
    doc.add_heading("2.2 STEP 5 in Detail: Evidence anchoring + Multi-LLM scoring → AUTO_ACCEPT | ESCALATE_LLM | NEEDS_REVIEW", level=2)
    doc.add_paragraph(
        "After structured data is parsed (STEP 4), the system (1) anchors each extracted field to evidence in the source text (evidence anchoring), "
        "and (2) computes a deterministic confidence and routes the run to one of four gates: AUTO_ACCEPT, AUTO_RETRY, ESCALATE_LLM, or NEEDS_REVIEW. "
        "Model self-reported confidence is not used; scoring is based on evidence coverage, concordance between candidates, extraction quality, "
        "rule pass rate, and invariant checks."
    )
    doc.add_paragraph("Evidence anchoring (what it does):", style="Heading 3")
    doc.add_paragraph(
        "Evidence anchoring (evidence_anchoring_service.anchor_evidence) takes the full extracted text and a list of fields (e.g. from the chosen "
        "candidate’s parsed_json: line_items and top-level totals). For each field it has a field_name and a value (e.g. account_name and amount). "
        "The service normalizes the value for search (e.g. strip, remove trailing .0 for numbers) and tries to find that value in the document text, "
        "trying literal, no-comma, and comma-formatted variants for numbers. When found, it records a snippet of text around the match (default ±80 "
        "characters) and, if per-page text is available, the page index. The result is a list of EvidenceEntry (field_name, page_index, snippet, bbox); "
        "coverage_pct is the fraction of fields for which a snippet was found. MasterJSONService.update_evidence_section writes these entries and "
        "coverage_pct into the Master JSON evidence section. This provides an audit trail: every extracted value can be traced to a location in the source."
    )
    doc.add_paragraph("Flattening candidate JSON for evidence:", style="Heading 3")
    doc.add_paragraph(
        "fields_from_candidate_parsed_json converts the chosen candidate’s parsed_json into a list of {field_name, value}. It iterates line_items "
        "(or units for rent_roll), using account_name/tenant_name/unit_id as field_name and amount/monthly_rent as value; then adds top-level keys "
        "such as total_assets, total_liabilities, total_equity, total_revenue, total_expenses, net_operating_income. That list is the input to anchor_evidence."
    )
    doc.add_paragraph("Multi-LLM scoring (deterministic confidence and gate):", style="Heading 3")
    doc.add_paragraph(
        "multi_llm_scoring.score_and_route loads the Master JSON and optionally the validation rule pass rate. It computes: (1) invariant_pass — "
        "for the document type (e.g. balance_sheet) it runs invariant checks (e.g. Assets = Liabilities + Equity within 0.01) on the first valid "
        "candidate’s parsed_json; (2) evidence_coverage from the Master JSON evidence section (or derived from fraction of evidence entries with snippets); "
        "(3) concordance — fraction of candidate pairs that agree on key totals (total_assets, total_liabilities, total_equity, total_revenue, total_expenses, "
        "net_operating_income); (4) extraction_score from the extraction section’s confidence_score (0–100 normalized to 0–1); (5) rule_score from "
        "rule_pass_rate if provided; (6) invariant_score (1.0 if invariant_pass, 0.0 if fail, 0.5 if not applicable)."
    )
    doc.add_paragraph(
        "Confidence is a weighted sum: 0.25×evidence_coverage + 0.25×concordance + 0.20×extraction_score + 0.15×rule_score + 0.15×invariant_score, "
        "clamped to [0, 1]. Extraction quality is considered low if extraction confidence_score < 50."
    )
    doc.add_paragraph("Gate routing rules:", style="Heading 3")
    doc.add_paragraph(
        "compute_gate(confidence, evidence_coverage, invariant_pass, extraction_quality_low) returns: NEEDS_REVIEW if invariant checks failed; "
        "AUTO_RETRY if extraction_quality_low; AUTO_ACCEPT if confidence ≥ 0.85 and evidence_coverage ≥ 0.90; NEEDS_REVIEW if confidence < 0.60 "
        "or evidence_coverage < 0.60; ESCALATE_LLM if 0.60 ≤ confidence < 0.85 or 0.60 ≤ evidence_coverage < 0.90 (and not already AUTO_ACCEPT); "
        "otherwise NEEDS_REVIEW. So: AUTO_ACCEPT = high confidence and strong evidence; ESCALATE_LLM = medium confidence or evidence (optional "
        "adversarial challenge can suggest corrections); NEEDS_REVIEW = low confidence/evidence or invariant failure; AUTO_RETRY = very low extraction quality."
    )
    doc.add_paragraph("Decision section and adversarial challenge:", style="Heading 3")
    doc.add_paragraph(
        "score_and_route builds a MasterJSONDecision: overall_gate, field_decisions (per line-item or key field: field_name, chosen_value, confidence, "
        "gate_outcome, rationale), synthesis_rationale (confidence, evidence_coverage, invariant_pass, and any fail_reasons). This is persisted via "
        "master_json_service.update_decision_section. When the gate is ESCALATE_LLM, the orchestrator may call the adversarial challenge service "
        "(run_challenge) with the low-confidence field list; the challenger LLM (e.g. Claude) suggests corrections; suggestions are appended to the "
        "Master JSON challenge_suggestions for human review."
    )
    doc.add_paragraph("Code references:", style="Heading 3")
    doc.add_paragraph("evidence_anchoring_service: anchor_evidence, fields_from_candidate_parsed_json, _find_snippet_in_text.", style="List Bullet")
    doc.add_paragraph("multi_llm_scoring: score_and_route, compute_confidence, compute_gate, _concordance_score, run_invariant_checks.", style="List Bullet")
    doc.add_paragraph("master_json_service: update_evidence_section, update_decision_section, append_challenge_suggestions.", style="List Bullet")
    doc.add_paragraph("adversarial_challenge_service: run_challenge (when gate == ESCALATE_LLM).", style="List Bullet")

    # --- 3. Tools in extraction engine ---
    doc.add_heading("3. Tools in the Extraction Engine", level=1)

    doc.add_heading("3.1 PDF Extraction Engines", level=2)
    engines = [
        ("PyMuPDF (fitz)", "Core", "Digital PDFs; fast text extraction. Library: PyMuPDF.", "90-95%"),
        ("PDFPlumber", "Core", "Tables and structure; layout-aware text.", "85-93%"),
        ("Camelot", "Optional", "Table extraction (stream/lattice).", "93-97%"),
        ("Tesseract OCR", "Optional", "Scanned/image-based PDFs. pytesseract + pdf2image.", "75-90%"),
        ("EasyOCR", "Optional", "OCR alternative; multi-language. EasyOCR library.", "75-90%"),
        ("LayoutLMv3", "Optional (ML)", "Microsoft LayoutLMv3; layout + token classification. Hugging Face.", "Varies"),
    ]
    p = doc.add_paragraph()
    p.add_run("Engine\tType\tPurpose\tTypical accuracy\n").bold = True
    for name, typ, purpose, acc in engines:
        doc.add_paragraph(f"{name}\t{typ}\t{purpose}\t{acc}")

    doc.add_heading("3.2 Classification & Validation (Pre-parse)", level=2)
    doc.add_paragraph("PDFClassifier: Classifies PDF as digital, scanned, mixed, table_heavy, form, or image_heavy. Uses PyMuPDF + PDFPlumber (table count) to choose extraction strategy.")
    doc.add_paragraph("QualityValidator: Runs 10 checks on extracted text: text length, special-character ratio, language consistency (langdetect), gibberish detection, word distribution, page consistency, empty pages, character distribution, whitespace ratio, confidence threshold. Produces confidence_score and overall_quality (excellent/good/acceptable/poor/failed).")

    doc.add_heading("3.3 Structured Parsing Tools", level=2)
    doc.add_paragraph("TemplateExtractor: Matches extracted text to ExtractionTemplate and ChartOfAccounts; uses fuzzy matching (fuzzywuzzy) to extract line items per account.")
    doc.add_paragraph("FinancialTableParser: Uses PDFPlumber to extract tables; parses balance sheet, income statement, cash flow, rent roll. Handles account codes (e.g. ####-####) and amount patterns; extracts header metadata (property_name, period_ending, etc.).")

    doc.add_heading("3.4 Confidence & Ensemble", level=2)
    doc.add_paragraph("ConfidenceEngine: Aggregates results from multiple engines with configurable weights (e.g. pymupdf 0.3, pdfplumber 0.4, camelot 0.3); detects conflicts and recommends resolution.")
    doc.add_paragraph("ModelScoringService: External scoring of engine outputs (no self-scoring by engines). Used by MultiEngineExtractor.")
    doc.add_paragraph("EnsembleEngine / EnhancedEnsembleEngine: Runs multiple engines and combines results (e.g. weighted voting, consensus).")

    doc.add_heading("3.5 Model Manager", level=2)
    doc.add_paragraph("ModelManager (singleton): Lazy-loads heavy engines (OCR, LayoutLM, EasyOCR, Camelot) once per worker to avoid repeated load times.")

    # --- 4. LLM and ML in extraction ---
    doc.add_heading("4. How LLM and ML Help in Data Extraction", level=1)

    doc.add_heading("4.1 LLM Usage", level=2)
    doc.add_paragraph("Local LLM (Ollama): Used for document classification and fallback extraction. Default model: deepseek-r1:14b (config: LLM_MODEL). Tasks: classify document type (balance_sheet, income_statement, rent_roll, cash_flow, etc.) and extract structured JSON when template/regex fails.")
    doc.add_paragraph("Multi-LLM extraction (AbeAI-style): When MULTI_LLM_EXTRACTION_ENABLED=True, 2–3 providers are called in parallel (Anthropic, OpenAI, Perplexity, Mistral, Gemini). Each returns a candidate JSON; results are stored in Master JSON. Scoring and gate routing (AUTO_ACCEPT, ESCALATE_LLM, NEEDS_REVIEW) decide final acceptance; adversarial challenge (Claude) can suggest corrections for low-confidence fields.")
    doc.add_paragraph("LLM adapters: anthropic (Claude), openai (GPT), perplexity, mistral, gemini. Config: ANTHROPIC_API_KEY, OPENAI_API_KEY, etc.")

    doc.add_heading("4.2 ML / Deep Learning", level=2)
    doc.add_paragraph("LayoutLMv3: Transformer model (microsoft/layoutlmv3-base) for document understanding; token classification and layout awareness. Used when LayoutLM engine is enabled (optional).")
    doc.add_paragraph("EasyOCR: Deep learning OCR (PyTorch). Used for scanned text when EasyOCR engine is loaded.")
    doc.add_paragraph("Tesseract OCR: Traditional OCR (not deep learning); used for scanned PDFs via pdf2image + pytesseract.")
    doc.add_paragraph("PyOD (optional): Used for anomaly detection (statistical/ML anomalies), not for raw text extraction.")

    doc.add_heading("4.3 Tool and Model Names Summary", level=2)
    models_list = [
        "Ollama (local): deepseek-r1:14b, qwen2.5:14b, llama3.2:3b, etc. (LLM_MODEL, OLLAMA_DEFAULT_MODEL)",
        "Anthropic: Claude (API)",
        "OpenAI: GPT (API)",
        "Perplexity / Mistral / Google: API models",
        "LayoutLMv3: microsoft/layoutlmv3-base (Hugging Face)",
        "EasyOCR: default language models (GPU/CPU)",
        "Tesseract: 5.5.0 (pytesseract)",
    ]
    for m in models_list:
        doc.add_paragraph(m, style="List Bullet")

    # --- 5. Data validation ---
    doc.add_heading("5. Data Validation (Post-Insert)", level=1)
    doc.add_paragraph("ValidationService runs business-logic rules per document type. Rules are stored in ValidationRule; each run creates a ValidationRun and per-rule ValidationResult.")
    doc.add_paragraph("Balance sheet: e.g. Assets = Liabilities + Equity (with tolerance), no negative totals where not allowed, header consistency.")
    doc.add_paragraph("Income statement: Revenue/expense logic, totals consistency.")
    doc.add_paragraph("Cash flow: Operating/Investing/Financing consistency, reconciliation checks.")
    doc.add_paragraph("Rent roll: Unit counts, rent totals, vacancy consistency.")
    doc.add_paragraph("Tolerance: configurable percentage (e.g. 1%) for rounding; failures can be error or warning.")

    doc.add_heading("5.1 Validation Flow Example", level=2)
    doc.add_paragraph("1. Upload completes extraction and inserts into balance_sheet_data.")
    doc.add_paragraph("2. validate_upload(upload_id) is called.")
    doc.add_paragraph("3. ValidationService loads active ValidationRules for that document_type.")
    doc.add_paragraph("4. For balance_sheet: _validate_balance_sheet() runs critical and warning checks (e.g. equation checks, sign checks).")
    doc.add_paragraph("5. Each check produces a ValidationResult (passed/failed, severity, message).")
    doc.add_paragraph("6. ValidationRun stores rules_version_hash, passed_count, failed_count; overall_passed = (failed_checks == 0).")

    # --- 6. Example end-to-end ---
    doc.add_heading("6. Example: Single Document Flow", level=1)
    doc.add_paragraph("1. User uploads 'Q4_Balance_Sheet.pdf'. File stored in MinIO; DocumentUpload created with document_type=balance_sheet.")
    doc.add_paragraph("2. Celery task extract_document(upload_id) runs. ExtractionOrchestrator.extract_and_parse_document(upload_id) is called.")
    doc.add_paragraph("3. PDF downloaded from MinIO. PDFClassifier classifies as 'digital'. MultiEngineExtractor.extract_with_validation() runs with strategy 'auto' (e.g. PyMuPDF or PDFPlumber selected). QualityValidator validates text; confidence_score 94%.")
    doc.add_paragraph("4. Optional: Multi-LLM generates 2–3 candidate JSONs; Master JSON updated; score_and_route() sets AUTO_ACCEPT.")
    doc.add_paragraph("5. _parse_and_insert_financial_data() uses TemplateExtractor and/or FinancialTableParser; line items inserted into balance_sheet_data.")
    doc.add_paragraph("6. ValidationService.validate_upload(upload_id) runs balance sheet rules; all pass.")
    doc.add_paragraph("7. Financial metrics calculated; anomaly detection runs; upload status set to completed; extraction log and Master JSON persisted.")

    doc.add_heading("7. Worked Examples", level=1)
    doc.add_heading("7.1 Extraction output example (balance sheet)", level=2)
    doc.add_paragraph(
        "After MultiEngineExtractor + FinancialTableParser, structured output looks like:"
    )
    doc.add_paragraph(
        '{"success": true, "document_type": "balance_sheet", "header": {"property_name": "Eastern Shore Plaza", '
        '"period_ending": "Dec 2023", "report_title": "Balance Sheet"}, '
        '"line_items": [{"account_code": "1000-1999", "account_name": "Cash and Equivalents", "amount": 125000.00, "line_type": "asset"}, ...], '
        '"total_items": 50, "extraction_method": "table"}'
    )
    doc.add_heading("7.2 Validation rule example", level=2)
    doc.add_paragraph(
        "Balance sheet critical rule: Total Assets must equal Total Liabilities + Equity within tolerance (e.g. 1%). "
        "ValidationService queries balance_sheet_data for the upload, sums assets and liabilities+equity, compares with tolerance; "
        "if outside tolerance, ValidationResult is created with passed=False, severity=error, message describing the mismatch."
    )
    doc.add_heading("7.3 Quality validation example", level=2)
    doc.add_paragraph(
        "QualityValidator runs after text extraction. Example: special-character ratio check. "
        "If extracted text has >30% non-alphanumeric characters (e.g. OCR artifacts), the check fails and contributes to lower confidence_score; "
        "overall_quality may become 'acceptable' or 'poor' and needs_review can be set."
    )

    doc.add_heading("8. References (Code Locations)", level=1)
    refs = [
        "Extraction orchestrator: backend/app/services/extraction_orchestrator.py",
        "Multi-engine extractor: backend/app/utils/extraction_engine.py",
        "Engines: backend/app/utils/engines/ (pymupdf_engine, pdfplumber_engine, camelot_engine, ocr_engine, easyocr_engine, layoutlm_engine)",
        "PDF classifier: backend/app/utils/pdf_classifier.py",
        "Quality validator: backend/app/utils/quality_validator.py",
        "Template extractor: backend/app/utils/template_extractor.py",
        "Financial table parser: backend/app/utils/financial_table_parser.py",
        "Confidence engine: backend/app/services/confidence_engine.py",
        "Validation service: backend/app/services/validation_service.py",
        "LLM extraction: backend/app/services/llm_extraction_service.py",
        "Local LLM: backend/app/services/local_llm_service.py",
        "Multi-LLM: backend/app/services/multi_llm_provider.py, llm_adapters/",
        "AI config: backend/app/core/config/ai.py",
        "EXTRACTION_SYSTEM_README: backend/EXTRACTION_SYSTEM_README.md",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph("Generated by REIMS2 backend/scripts/generate_extraction_engine_doc.py")

    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return 0

if __name__ == "__main__":
    exit(main())
