#!/usr/bin/env python3
"""
Generate Natural Language Query (NLQ) User Manual

Creates a comprehensive .docx user manual for the Natural Language Query feature,
including purpose, how to use, examples, and end-user guidance.

Output: /home/hsthind/REIMS - Documents/REIMS2_Natural_Language_Query_User_Manual.docx
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Use REAL screenshots from this folder. See SCREENSHOT_CAPTURE_GUIDE.md for capture instructions.
SCREENSHOTS_DIR = Path("/home/hsthind/REIMS - Documents/reims-screenshots")


def add_screenshot(doc, filename, caption, width=5.5):
    """Add real screenshot if it exists; otherwise add placeholder with capture instructions."""
    path = SCREENSHOTS_DIR / filename
    if path.exists():
        try:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(path), width=Inches(width))
            if caption:
                cap = doc.add_paragraph()
                cap.add_run(caption).italic = True
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            doc.add_paragraph(f"[Insert screenshot: {filename}]")
    else:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("⚠️ INSERT REAL SCREENSHOT: ").bold = True
        p.add_run(f"Capture from the live REIMS app. Save as '{filename}' in reims-screenshots folder. See SCREENSHOT_CAPTURE_GUIDE.md for exact steps.")
        p.paragraph_format.space_before = Pt(6)
    doc.add_paragraph()


def main():
    doc = Document()
    doc.add_heading("REIMS2 Natural Language Query", 0)
    doc.add_paragraph()
    doc.add_heading("User Manual", level=1)
    doc.add_paragraph(
        "This manual explains how to use the Natural Language Query (NLQ) feature to ask questions about "
        "your financial data in plain English. The AI-powered system understands temporal expressions, "
        "financial formulas, and data queries—no SQL or technical knowledge required."
    )
    doc.add_paragraph()

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "Purpose of Natural Language Query",
        "How to Access Natural Language Query",
        "NLQ Page Layout",
        "Step-by-Step: How to Ask a Question",
        "Supported Query Types",
        "Temporal Expressions (Dates & Periods)",
        "Example Queries by Category",
        "Understanding the Response",
        "Property Filter",
        "Formula Browser",
        "Tips for Best Results",
        "Troubleshooting",
    ]
    for item in toc_items:
        doc.add_paragraph(f"• {item}", style="List Bullet")
    doc.add_paragraph()

    # Purpose
    doc.add_heading("Purpose of Natural Language Query", level=1)
    doc.add_paragraph(
        "Natural Language Query lets you ask questions about your REIMS financial data in plain English. "
        "Instead of writing SQL or navigating multiple reports, you can ask:"
    )
    bullets = [
        "Financial data questions (cash, revenue, expenses, NOI, occupancy, DSCR, etc.)",
        "Formula and calculation questions (e.g., How is DSCR calculated?)",
        "Temporal questions (last 3 months, Q4 2025, November 2025, YTD)",
        "Comparison queries (YTD vs last year, Q4 2025 vs Q4 2024)",
        "Audit and history queries (who changed what, when)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()

    # How to Access
    doc.add_heading("How to Access Natural Language Query", level=1)
    doc.add_paragraph(
        "Navigate to the AI Assistant (or Natural Language Query) from the left sidebar. "
        "You can also use the shortcut Ctrl/Cmd + 7, or go to #nlq-search. "
        "The NLQ interface is also available within the Financials page for context-specific queries."
    )
    doc.add_paragraph()

    # Layout
    doc.add_heading("NLQ Page Layout", level=1)
    add_screenshot(doc, "nlq-main-page.png", "Figure 1: Natural Language Query main page")
    doc.add_paragraph()

    doc.add_heading("Page Components", level=2)
    comps = [
        ("NLQ System Online", "Green status banner showing the system is ready. Displays agent count and features."),
        ("Filter by Property (Optional)", "Dropdown to restrict queries to a specific property (e.g., Eastern Shore Plaza). Leave as 'All Properties' for portfolio-wide queries."),
        ("Ask a Question", "Main search input. Type your question in plain English and click Ask or press Ctrl/Cmd+Enter."),
        ("Quick Questions", "Click any suggested question to populate the input. Great for first-time users."),
        ("Example Queries", "Expandable sections (Financial Data, Formulas & Calculations, Temporal Queries, Audit & History) with clickable examples."),
        ("Supported Features", "Badges showing: Temporal Queries (10+ types), 50+ Financial Formulas, Multi-Statement Analysis, Audit Trail, Reconciliation, Natural Language."),
        ("Formula Browser", "Browse 50+ financial formulas by category. Expand any formula to see its definition and explanation."),
    ]
    for name, desc in comps:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)
    doc.add_paragraph()

    # Step by Step
    doc.add_heading("Step-by-Step: How to Ask a Question", level=1)
    steps = [
        "Open the Natural Language Query page (#nlq-search).",
        "Optional: Select a property from the Filter by Property dropdown if you want results for a specific property only.",
        "Type your question in the search box. For example: 'What was cash position in November 2025?'",
        "Click Ask (or press Ctrl/Cmd + Enter).",
        "Wait a few seconds while the system analyzes your question and retrieves data. You will see 'Analyzing your question...' during processing.",
        "Read the answer. The response includes confidence score and execution time. Expand 'View Raw Data' or 'Query Details' for more context.",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}")
    doc.add_paragraph()

    # Supported Query Types
    doc.add_heading("Supported Query Types", level=1)
    doc.add_paragraph("The NLQ system supports these categories of questions:")
    doc.add_paragraph()
    doc.add_heading("1. Financial Data Queries", level=2)
    doc.add_paragraph("Ask about balance sheet, income statement, cash flow, metrics, and rent roll data.")
    doc.add_paragraph("Examples:")
    for ex in [
        "What was the cash position in November 2025?",
        "Show me total revenue for Q4 2025",
        "What are total assets for property ESP?",
        "Show operating expenses for last month",
        "Compare net income YTD vs last year",
    ]:
        doc.add_paragraph(f"  • {ex}")
    doc.add_paragraph()

    doc.add_heading("2. Formulas & Calculations", level=2)
    doc.add_paragraph("Ask how a metric is calculated or request a calculated value.")
    doc.add_paragraph("Examples:")
    for ex in [
        "How is DSCR calculated?",
        "What is the formula for Current Ratio?",
        "Explain NOI calculation",
        "Calculate DSCR for property ESP in November 2025",
        "What is the benchmark for good DSCR?",
    ]:
        doc.add_paragraph(f"  • {ex}")
    doc.add_paragraph()

    doc.add_heading("3. Temporal Queries", level=2)
    doc.add_paragraph("Use natural language dates and periods. The system understands many formats.")
    doc.add_paragraph("Examples:")
    for ex in [
        "Show data for last 3 months",
        "Compare Q4 2025 vs Q4 2024",
        "Year to date revenue",
        "Month to date expenses",
        "Between August and December 2025",
    ]:
        doc.add_paragraph(f"  • {ex}")
    doc.add_paragraph()

    doc.add_heading("4. Audit & History", level=2)
    doc.add_paragraph("Query audit trails and change history.")
    doc.add_paragraph("Examples:")
    for ex in [
        "Who changed cash position in November 2025?",
        "Show me audit history for property ESP",
        "What was modified last week?",
        "List all changes by user John Doe",
    ]:
        doc.add_paragraph(f"  • {ex}")
    doc.add_paragraph()

    # Temporal Expressions
    doc.add_heading("Temporal Expressions (Dates & Periods)", level=1)
    doc.add_paragraph("You can use many natural ways to specify time. The system supports:")
    temp_table = doc.add_table(rows=1, cols=2)
    temp_table.style = "Table Grid"
    hdr = temp_table.rows[0].cells
    hdr[0].text = "Expression Type"
    hdr[1].text = "Examples"
    for row in [
        ("Absolute (Month + Year)", "November 2025, Jan 2026, December 2025"),
        ("ISO Date", "2025-11-15, 2025-11-01"),
        ("Year Only", "in 2025, for 2025"),
        ("Relative Periods", "last 3 months, last year, previous quarter"),
        ("Fiscal Periods", "Q4 2025, Q1 2026, fiscal year 2025"),
        ("Keywords", "YTD (year-to-date), MTD (month-to-date), QTD (quarter-to-date)"),
        ("Date Ranges", "between August and December 2025, from Jan to Mar 2025"),
        ("Natural Month Names", "August, September, October (with year context)"),
    ]:
        r = temp_table.add_row()
        r.cells[0].text = row[0]
        r.cells[1].text = row[1]
    doc.add_paragraph()

    # Example Queries Table
    doc.add_heading("Example Queries by Category", level=1)
    ex_table = doc.add_table(rows=1, cols=3)
    ex_table.style = "Table Grid"
    eh = ex_table.rows[0].cells
    eh[0].text = "Question"
    eh[1].text = "What It Returns"
    eh[2].text = "Use Case"
    examples = [
        ("What was cash position in November 2025?", "Total cash for that month", "Liquidity check"),
        ("Show total revenue for Q4 2025", "Revenue for Oct–Dec 2025", "Quarterly review"),
        ("How is DSCR calculated?", "Formula, inputs, explanation", "Learning or validation"),
        ("Calculate Current Ratio for last month", "Current Ratio value", "Liquidity metric"),
        ("Compare net income YTD vs last year", "Side-by-side comparison", "Performance analysis"),
        ("What are total assets for ESP?", "Assets for that property", "Balance sheet review"),
        ("Show occupancy for last 3 months", "Occupancy trend", "Operational monitoring"),
        ("Which properties have highest operating expense ratio?", "Ranked list", "Portfolio analysis"),
    ]
    for q, ret, use in examples:
        r = ex_table.add_row()
        r.cells[0].text = q
        r.cells[1].text = ret
        r.cells[2].text = use
    doc.add_paragraph()

    # Understanding the Response
    doc.add_heading("Understanding the Response", level=1)
    add_screenshot(doc, "nlq-query-result.png", "Figure 2: Example query and response")
    doc.add_paragraph()
    doc.add_paragraph("The response card includes:")
    resp_items = [
        ("Answer", "Natural language answer to your question, often with specific numbers and context."),
        ("Confidence", "0–100%. Green (80%+) = high confidence; Yellow (60–79%) = moderate; Red (<60%) = low. Use this to gauge reliability."),
        ("Execution Time", "How long the query took (e.g., 1.2s). Cached queries are faster (<100ms)."),
        ("View Raw Data", "Expandable section showing the underlying data records in JSON format."),
        ("Query Details", "Expandable section with time period (if temporal) and agents used."),
    ]
    for name, desc in resp_items:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)
    doc.add_paragraph()

    # Property Filter
    doc.add_heading("Property Filter", level=1)
    doc.add_paragraph(
        "Use 'Filter by Property' to scope queries to a single property. When a property is selected, "
        "the system will focus answers on that property (e.g., Eastern Shore Plaza). Leave as 'All Properties' "
        "for portfolio-wide or comparison queries. The filter is optional."
    )
    doc.add_paragraph()

    # Formula Browser
    doc.add_heading("Formula Browser", level=1)
    doc.add_paragraph(
        "The Formula Browser lists 50+ financial formulas by category: Liquidity, Leverage, Mortgage, "
        "Income Statement, Rent Roll, and more. Select a category from the dropdown, then expand any formula "
        "to see its definition, formula expression, and explanation. Useful for understanding metrics before asking calculation questions."
    )
    doc.add_paragraph()

    # Tips
    doc.add_heading("Tips for Best Results", level=1)
    tips = [
        "Be specific: 'What was cash in November 2025?' is better than 'cash?'",
        "Include time when relevant: 'revenue for Q4 2025' vs 'revenue'",
        "Use property code if needed: 'DSCR for ESP' when you have multiple properties",
        "Start with example queries: Click a suggested question to see the format",
        "Use natural language: 'last 3 months' works; you don't need exact dates",
        "Rephrase if unclear: If the answer seems off, try a different wording",
    ]
    for t in tips:
        doc.add_paragraph(t, style="List Bullet")
    doc.add_paragraph()

    # Troubleshooting
    doc.add_heading("Troubleshooting", level=1)
    doc.add_paragraph()
    doc.add_heading("No answer or error message", level=2)
    doc.add_paragraph(
        "Ensure the backend is running and NLQ is configured. Check 'NLQ System Online' status. "
        "If you see an error, try rephrasing your question or verifying that the required data exists for the selected property and period."
    )
    doc.add_heading("Low confidence score", level=2)
    doc.add_paragraph(
        "A low confidence score may mean ambiguous question or missing data. Be more specific, include a time period, or select a property."
    )
    doc.add_heading("Wrong property or period", level=2)
    doc.add_paragraph(
        "Use the Property Filter dropdown to restrict to a specific property. Include the period explicitly in your question (e.g., November 2025)."
    )
    doc.add_paragraph()

    # Maintenance
    doc.add_heading("Maintenance", level=1)
    doc.add_paragraph("Last Updated: January 31, 2026")
    doc.add_paragraph("Document Version: 1.0")

    out_dir = Path("/home/hsthind/REIMS - Documents")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "REIMS2_Natural_Language_Query_User_Manual.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return str(out_path)


if __name__ == "__main__":
    main()
