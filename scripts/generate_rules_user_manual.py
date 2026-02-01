#!/usr/bin/env python3
"""
Generate comprehensive REIMS2 Rules End-User Manual as .docx

Creates a detailed user manual documenting all reconciliation rules with:
- Rule ID, Name, Formula
- Document types included
- Accounts and line items
- Calculations performed
- Examples and validation guidance

Output: /home/hsthind/REIMS - Documents/REIMS2_Rules_End_User_Manual.docx
"""
import ast
import json
import re
import sys
from pathlib import Path

# Add backend to path for imports
script_dir = Path(__file__).resolve().parent
repo_root = script_dir.parent
backend_path = repo_root / "backend"
sys.path.insert(0, str(repo_root))
sys.path.insert(0, str(backend_path))

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

# Rule enrichment data - accounts, line items, document mappings
RULE_ENRICHMENT = {
    "BS-1": {
        "formula_override": "Total Assets - (Total Liabilities & Capital)",
        "accounts": ["TOTAL ASSETS", "TOTAL LIABILITIES & CAPITAL (Total Liabilities + Total Capital)"],
        "line_items": ["1999-0000 Total Assets", "2999-0000 Total Liabilities", "3999-0000 Total Capital"],
        "calc_detail": "Subtract (Total Liabilities + Total Equity) from Total Assets. Result must be zero (within threshold).",
        "example": "Source Value (Actual): $23,899,054.41 | Target Value (Expected): $23,899,054.41 | Threshold: 0.01 — PASS when variance ≤ $0.01",
    },
    "BS-2": {"accounts": ["0122-0000 Cash Operating"], "line_items": ["Cash - Operating"], "calc_detail": "Compare to baseline $3,375.45.", "example": ""},
    "BS-3": {"accounts": ["Total Current Assets"], "line_items": ["0499-9000"], "calc_detail": "Verify sum > 0.", "example": ""},
    "BS-4": {"accounts": ["Total Current Assets", "Total Current Liabilities"], "line_items": ["0499-9000", "2499-9000"], "calc_detail": "CA / CL >= 1.0", "example": ""},
    "BS-5": {"accounts": ["Total Current Assets", "Total Current Liabilities"], "line_items": ["0499-9000", "2499-9000"], "calc_detail": "CA - CL > 0", "example": ""},
    "BS-6": {"accounts": ["0510-0000 Land"], "line_items": ["Land"], "calc_detail": "Land > 0", "example": ""},
    "BS-7": {"accounts": ["1230-0000 Accumulated Depreciation"], "line_items": ["Accum Depr"], "calc_detail": "Current period abs(value) >= Prior period abs(value)", "example": ""},
    "BS-8": {"accounts": ["Accum Depr Buildings"], "line_items": ["1061-0000"], "calc_detail": "Monthly increase check", "example": ""},
    "BS-9": {"accounts": ["Total Liabilities", "TOTAL ASSETS"], "line_items": ["2999-0000", "1999-0000"], "calc_detail": "Liabilities / Assets <= 0.85", "example": ""},
    "BS-10": {"accounts": ["5 Year Improvements"], "line_items": ["0710-0000"], "calc_detail": "Baseline comparison -$1,025,187", "example": ""},
    "BS-11": {"accounts": ["TI/Current Improvements"], "line_items": ["0950-0000"], "calc_detail": ">= 0", "example": ""},
    "BS-12": {"accounts": ["30 Year Roof"], "line_items": ["0815-0000"], "calc_detail": "> 0", "example": ""},
    "BS-13": {"accounts": ["30 Year HVAC"], "line_items": ["0816-0000"], "calc_detail": ">= 0", "example": ""},
    "BS-14": {"accounts": ["Deposits"], "line_items": ["1210-0000"], "calc_detail": "Baseline $20,900", "example": ""},
    "BS-15": {"accounts": ["Loan Costs"], "line_items": ["1920-0000"], "calc_detail": "Baseline $268,752.01", "example": ""},
    "BS-16": {"accounts": ["Accum Amort Loan Costs"], "line_items": ["Loan Cost Amortization"], "calc_detail": "Non-decreasing (abs)", "example": ""},
    "BS-17": {"accounts": ["Accum Amort Other"], "line_items": ["Other Amortization"], "calc_detail": "Baseline -$36,621.19", "example": ""},
    "BS-18": {"accounts": ["External Lease Commissions"], "line_items": ["Ext Lease Comm"], "calc_detail": "> 0", "example": ""},
    "BS-19": {"accounts": ["Internal Lease Commissions"], "line_items": ["Int Lease Comm"], "calc_detail": "> 0", "example": ""},
    "BS-20": {"accounts": ["Prepaid Insurance"], "line_items": ["Prepaid Ins"], "calc_detail": "> 0", "example": ""},
    "BS-21": {"accounts": ["Prepaid Expenses"], "line_items": ["Prepaid Exp"], "calc_detail": ">= 0", "example": ""},
    "BS-22": {"accounts": ["A/P 5Rivers"], "line_items": ["AP 5Rivers"], "calc_detail": "Baseline $31,683.54", "example": ""},
    "BS-23": {"accounts": ["A/P Eastchase"], "line_items": ["AP Eastchase"], "calc_detail": "Baseline $354.54", "example": ""},
    "BS-24": {"accounts": ["Loans Payable 5Rivers"], "line_items": ["Loans 5Rivers"], "calc_detail": "Baseline $1,810,819.58", "example": ""},
    "BS-25": {"accounts": ["Deposit Refundable"], "line_items": ["Deposit Refund"], "calc_detail": "Baseline $49,791.31", "example": ""},
    "BS-26": {"accounts": ["Accrued Expenses"], "line_items": [], "calc_detail": "Tracking", "example": ""},
    "BS-27": {"accounts": ["A/P Trade"], "line_items": [], "calc_detail": "Tracking", "example": ""},
    "BS-28": {"accounts": ["Property Tax Payable"], "line_items": [], "calc_detail": "Accumulation tracking", "example": ""},
    "BS-29": {"accounts": ["Rent In Advance"], "line_items": [], "calc_detail": "Tracking", "example": ""},
    "BS-30": {"accounts": ["Partners Contribution"], "line_items": ["Partners Contrib"], "calc_detail": "Baseline $5,684,514.69", "example": ""},
    "BS-31": {"accounts": ["Beginning Equity"], "line_items": ["Beg Equity"], "calc_detail": "Baseline $1,786,413.82", "example": ""},
    "BS-32": {"accounts": ["Distributions"], "line_items": [], "calc_detail": "<= 0", "example": ""},
    "BS-33": {"accounts": ["Current Period Earnings", "Net Income"], "line_items": ["Curr Earnings", "IS Net Income"], "calc_detail": "Delta BS Earnings = IS Net Income", "example": ""},
    "BS-34": {"accounts": ["Total Capital"], "line_items": ["3999-0000"], "calc_detail": "Sum verification", "example": ""},
    "BS-35": {"accounts": ["Total Capital"], "line_items": [], "calc_detail": "Period-over-period tracking", "example": ""},
    # Cash Flow
    "CF-1": {"accounts": ["Operating", "Investing", "Financing", "Net Change"], "line_items": ["Category totals"], "calc_detail": "Operating + Investing + Financing = Net Change in Cash", "example": ""},
    "CF-2": {"accounts": ["Beginning Cash", "Net Change", "Ending Cash"], "line_items": ["Cash reconciliation"], "calc_detail": "Beginning + Net Change = Ending Cash", "example": ""},
    "CF-6": {"accounts": ["Net Income", "Adjustments"], "line_items": ["CF Operating section"], "calc_detail": "Net Income + Total Adjustments = Operating Cash Flow", "example": ""},
    # Income Statement
    "IS-1": {"accounts": ["Net Income", "NOI", "Interest", "Depreciation", "Amortization"], "line_items": ["NET INCOME", "NET OPERATING INCOME"], "calc_detail": "NI = NOI - (Interest + Depr + Amort)", "example": ""},
    "IS-NOI": {"accounts": ["Total Income", "Operating Expenses", "NOI"], "line_items": ["TOTAL INCOME", "NET OPERATING INCOME"], "calc_detail": "NOI = Revenue - Operating Expenses", "example": ""},
    # Three Statement
    "3S-1": {"accounts": ["BS Cash", "CF Ending Cash"], "line_items": ["Balance Sheet Cash", "Cash Flow Ending"], "calc_detail": "BS Cash balance = CF Ending Cash", "example": ""},
    "3S-3": {"accounts": ["IS Net Income", "BS Current Period Earnings"], "line_items": ["Net Income", "Current Period Earnings"], "calc_detail": "IS Net Income = Change in BS Earnings", "example": ""},
    "3S-4": {"accounts": ["IS Net Income", "CF Net Income line"], "line_items": [], "calc_detail": "IS NI = CF starting Net Income", "example": ""},
}

DOC_TYPE_LABELS = {
    "balance_sheet": "Balance Sheet",
    "income_statement": "Income Statement",
    "cash_flow": "Cash Flow Statement",
    "mortgage_statement": "Mortgage Statement",
    "rent_roll": "Rent Roll",
    "all": "All Document Types",
}


def parse_seed_rules():
    """Parse rules from forensic_calculated_rules_seed.py"""
    seed_path = backend_path / "app" / "db" / "seeds" / "forensic_calculated_rules_seed.py"
    content = seed_path.read_text()
    rules = []
    blocks = re.split(r"# =+ [A-Z ]+ =+", content)
    for block in blocks:
        for line in block.split("\n"):
            line = line.strip()
            if '{"rule_id"' not in line or '"rule_name"' not in line:
                continue
            try:
                line_clean = line.rstrip(",").strip()
                data = ast.literal_eval(line_clean)
                rules.append({
                    "rule_id": data["rule_id"],
                    "rule_name": data["rule_name"],
                    "formula": data["formula"],
                    "doc_scope": data["doc_scope"],
                    "severity": data["severity"],
                })
            except (ValueError, SyntaxError, KeyError):
                pass
    return rules


def scrape_rule_names_from_python():
    """Extract rule_id -> rule_name from Python rule mixins."""
    rules_dir = backend_path / "app" / "services" / "rules"
    mapping = {}
    pattern = re.compile(r'rule_id\s*=\s*["\']([^"\']+)["\'][^}]*?rule_name\s*=\s*["\']([^"\']+)["\']', re.DOTALL)
    for py_file in rules_dir.glob("*.py"):
        if py_file.name == "rule_registry.py":
            continue
        content = py_file.read_text()
        for m in pattern.finditer(content):
            rid, rname = m.group(1), m.group(2)
            if rid not in mapping or len(rname) > len(mapping.get(rid, "")):
                mapping[rid] = rname
    return mapping


def get_all_rules_from_registry():
    """Get all rule IDs from rule_registry and enrich with names from Python + seed."""
    sys.path.insert(0, str(backend_path))
    from app.services.rules.rule_registry import ALL_RULE_IDS

    seed_rules = {r["rule_id"]: r for r in parse_seed_rules()}
    scraped_names = scrape_rule_names_from_python()

    # Doc scope by prefix (longer prefixes first)
    scope_map = [
        ("FA-CASH-", ["cash_flow"]),
        ("FA-MORT-", ["mortgage_statement"]),
        ("FA-RR-", ["rent_roll"]),
        ("FA-WC-", ["balance_sheet", "cash_flow"]),
        ("FA-", ["balance_sheet", "cash_flow", "income_statement"]),
        ("BS-", ["balance_sheet"]),
        ("CF-", ["cash_flow"]),
        ("IS-", ["income_statement"]),
        ("MST-", ["mortgage_statement"]),
        ("RRBS-", ["rent_roll", "balance_sheet"]),
        ("RR-", ["rent_roll"]),
        ("3S-", ["balance_sheet", "cash_flow", "income_statement"]),
        ("AUDIT-", ["balance_sheet", "income_statement", "cash_flow", "mortgage_statement", "rent_roll"]),
        ("DQ-", ["balance_sheet", "income_statement", "cash_flow"]),
        ("ANALYTICS-", ["all"]),
        ("COVENANT-", ["all"]),
        ("BENCHMARK-", ["all"]),
        ("TREND-", ["all"]),
        ("STRESS-", ["all"]),
        ("DASHBOARD-", ["all"]),
        ("WCR-", ["balance_sheet", "cash_flow"]),
        ("MCI-", ["mortgage_statement", "balance_sheet", "cash_flow", "income_statement"]),
    ]

    def infer_scope(rid):
        for prefix, scope in scope_map:
            if rid.startswith(prefix):
                return scope
        return ["all"]

    rules = []
    for rid in sorted(ALL_RULE_IDS, key=natural_sort_key):
        if rid in seed_rules:
            rules.append(seed_rules[rid])
        else:
            name = scraped_names.get(rid, rid.replace("-", " ").title())
            rules.append({
                "rule_id": rid,
                "rule_name": name,
                "formula": "See implementation",
                "doc_scope": infer_scope(rid),
                "severity": "medium",
            })
    return rules


def add_rule_section(doc, rule, enrichment):
    """Add a detailed rule section to the document"""
    # Rule heading
    h = doc.add_heading(f"{rule['rule_id']}: {rule['rule_name']}", level=2)
    h.runs[0].bold = True

    # Table for rule details
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    def add_row(label, value):
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value) if value else "—"
        row.cells[0].paragraphs[0].runs[0].bold = True

    add_row("Rule Type", "Calculated Rule")
    # Use display formula from enrichment if available (e.g., BS-1 user-friendly formula)
    formula_display = (enrichment or {}).get("formula_override")
    add_row("Formula Definition", formula_display or rule.get("formula", ""))
    add_row("Severity", rule.get("severity", "").upper())

    doc_scopes = rule.get("doc_scope", [])
    doc_labels = [DOC_TYPE_LABELS.get(d, d) for d in doc_scopes]
    add_row("Document Types", ", ".join(doc_labels) if doc_labels else "—")

    if enrichment:
        add_row("Accounts Included", ", ".join(enrichment.get("accounts", [])) or "—")
        add_row("Line Items (Account Codes)", ", ".join(enrichment.get("line_items", [])) or "—")
        add_row("Calculation Performed", enrichment.get("calc_detail", "") or "—")
        if enrichment.get("example"):
            add_row("Example", enrichment.get("example", "") or "—")
    else:
        add_row("Accounts Included", "See formula for account references")
        add_row("Calculation Performed", rule.get("formula", ""))

    # Threshold note
    p = doc.add_paragraph()
    p.add_run("Threshold: ").bold = True
    p.add_run("Default tolerance is 0.01 for calculated rules. Rules may use absolute ($) or percentage (%) tolerance. Check rule configuration for specific thresholds.")
    p.paragraph_format.space_before = Pt(6)
    doc.add_paragraph()


def main():
    rules = get_all_rules_from_registry()
    print(f"Loaded {len(rules)} rules (full coverage)")

    doc = Document()
    # Title
    title = doc.add_heading("REIMS2 Reconciliation Rules", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("End-User Validation Manual", level=1)
    doc.add_paragraph(
        "This manual provides detailed information for validating all reconciliation rules in the REIMS2 Financial Integrity system. "
        "Use this document to understand what each rule checks, which accounts and documents are involved, and how to interpret results."
    )

    # Rule count explanation
    doc.add_heading("Understanding the Rule Count (377 Active Rules)", level=2)
    doc.add_paragraph(
        "The REIMS2 system displays \"377 Rules Active\" (or similar) for a given property and period. This count represents the number of rule results that executed for your specific context. "
        "The system implements 321 unique rule definitions in code. The higher count (377) occurs when:\n\n"
        "• Dynamic sub-rules are generated (e.g., AUDIT-5 creates separate Tax, Insurance, Reserve checks).\n"
        "• Rules run per-account or per-escrow-type.\n"
        "• Multiple cross-document validations produce distinct results.\n\n"
        "The count can also be lower when:\n"
        "• Rules require specific document types (Balance Sheet, Income Statement, Cash Flow, Mortgage Statement, Rent Roll) that are missing.\n"
        "• Period-over-period rules need prior period data.\n\n"
        f"This manual documents all {len(rules)} rule definitions with full detail for core financial rules and reference tables for Audit, Data Quality, Analytics, and other categories."
    )
    doc.add_paragraph()

    doc.add_heading("Table of Contents", level=2)
    for item in ["Understanding the Rule Count", "How to Use This Manual", "Rule Validation Workflow",
                 "Balance Sheet Rules (BS)", "Income Statement Rules (IS)", "Cash Flow Rules (CF)",
                 "Mortgage Statement Rules (MST)", "Rent Roll Rules (RR)", "Three-Statement Rules (3S)",
                 "Audit Rules (AUDIT)", "Data Quality Rules (DQ)", "Analytics Rules (ANALYTICS)",
                 "Covenant, Benchmark, Trend, Stress, Dashboard Rules", "Forensic & Cross-Document Rules",
                 "Rent Roll–Balance Sheet Rules (RRBS)", "Rule Count Summary", "Document Types Reference", "Maintenance Information"]:
        doc.add_paragraph(f"• {item}", style="List Bullet")
    doc.add_paragraph()

    doc.add_heading("How to Use This Manual", level=2)
    doc.add_paragraph(
        "1. Navigate to the Financial Integrity Hub in REIMS2.\n"
        "2. Select your property and reporting period.\n"
        "3. Review the Overall Status and rule counts.\n"
        "4. Click on any rule card to see current evaluation values.\n"
        "5. Use this manual to understand the rule's formula, accounts, and expected behavior.\n"
        "6. For failing rules, verify source documents and line item extractions."
    )
    doc.add_paragraph()

    doc.add_heading("Rule Validation Workflow", level=2)
    doc.add_paragraph(
        "• PASS: Rule met its validation criteria (within threshold).\n"
        "• FAIL/Variance: Rule did not meet criteria—investigate source data.\n"
        "• SKIP: Rule did not run (missing data or not applicable).\n"
        "• INFO: Informational result—no action required unless noted."
    )
    doc.add_paragraph()

    # Group rules by prefix
    groups = {}
    for r in rules:
        rid = r["rule_id"]
        if rid.startswith("BS-"):
            groups.setdefault("Balance Sheet Rules (BS)", []).append(r)
        elif rid.startswith("CF-"):
            groups.setdefault("Cash Flow Rules (CF)", []).append(r)
        elif rid.startswith("IS-"):
            groups.setdefault("Income Statement Rules (IS)", []).append(r)
        elif rid.startswith("MST-"):
            groups.setdefault("Mortgage Statement Rules (MST)", []).append(r)
        elif rid.startswith("RR-") and not rid.startswith("RRBS-"):
            groups.setdefault("Rent Roll Rules (RR)", []).append(r)
        elif rid.startswith("3S-"):
            groups.setdefault("Three-Statement Rules (3S)", []).append(r)
        elif rid.startswith("AUDIT-"):
            groups.setdefault("Audit Rules (AUDIT)", []).append(r)
        elif rid.startswith("DQ-"):
            groups.setdefault("Data Quality Rules (DQ)", []).append(r)
        elif rid.startswith("ANALYTICS-"):
            groups.setdefault("Analytics Rules (ANALYTICS)", []).append(r)
        elif rid.startswith("COVENANT-"):
            groups.setdefault("Covenant Rules (COVENANT)", []).append(r)
        elif rid.startswith("BENCHMARK-"):
            groups.setdefault("Benchmark Rules (BENCHMARK)", []).append(r)
        elif rid.startswith("TREND-"):
            groups.setdefault("Trend Rules (TREND)", []).append(r)
        elif rid.startswith("STRESS-"):
            groups.setdefault("Stress Rules (STRESS)", []).append(r)
        elif rid.startswith("DASHBOARD-"):
            groups.setdefault("Dashboard Rules (DASHBOARD)", []).append(r)
        elif rid.startswith("FA-") or rid.startswith("WCR-") or rid.startswith("MCI-"):
            groups.setdefault("Forensic & Cross-Document Rules", []).append(r)
        elif rid.startswith("RRBS-"):
            groups.setdefault("Rent Roll–Balance Sheet Rules (RRBS)", []).append(r)
        else:
            groups.setdefault("Other Rules", []).append(r)

    # Categories with full detail vs. condensed table
    full_detail_cats = ["Balance Sheet Rules (BS)", "Income Statement Rules (IS)", "Cash Flow Rules (CF)",
                        "Mortgage Statement Rules (MST)", "Rent Roll Rules (RR)", "Three-Statement Rules (3S)"]
    group_order = full_detail_cats + [
        "Audit Rules (AUDIT)", "Data Quality Rules (DQ)", "Analytics Rules (ANALYTICS)",
        "Covenant Rules (COVENANT)", "Benchmark Rules (BENCHMARK)", "Trend Rules (TREND)",
        "Stress Rules (STRESS)", "Dashboard Rules (DASHBOARD)",
        "Forensic & Cross-Document Rules", "Rent Roll–Balance Sheet Rules (RRBS)", "Other Rules",
    ]

    for gname in group_order:
        if gname not in groups:
            continue
        doc.add_heading(gname, level=1)
        group_rules = sorted(groups[gname], key=lambda x: natural_sort_key(x["rule_id"]))
        if gname == "Balance Sheet Rules (BS)":
            doc.add_heading("Featured Example: BS-1 Accounting Equation", level=2)
            doc.add_paragraph(
                "The Accounting Equation is the most critical rule. It validates that the fundamental equation holds: "
                "Total Assets = Total Liabilities & Capital (i.e., Total Liabilities + Total Equity)."
            )
            ex_table = doc.add_table(rows=5, cols=2)
            ex_table.style = "Table Grid"
            ex_rows = [
                ("Source Value (Actual)", "$23,899,054.41 — Total Assets from Balance Sheet"),
                ("Target Value (Expected)", "$23,899,054.41 — Total Liabilities & Capital"),
                ("Formula", "Total Assets - (Total Liabilities & Capital) = 0"),
                ("Threshold", "0.01 — Variance within $0.01 is acceptable"),
                ("Status", "PASS when both values match within threshold"),
            ]
            for i, (k, v) in enumerate(ex_rows):
                ex_table.rows[i].cells[0].text = k
                ex_table.rows[i].cells[1].text = v
                ex_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            doc.add_paragraph()
        if gname in full_detail_cats:
            for rule in group_rules:
                enrichment = RULE_ENRICHMENT.get(rule["rule_id"], {})
                add_rule_section(doc, rule, enrichment)
        else:
            # Condensed table for other categories
            tbl = doc.add_table(rows=0, cols=3)
            tbl.style = "Table Grid"
            rh = tbl.add_row()
            rh.cells[0].text = "Rule ID"
            rh.cells[1].text = "Rule Name"
            rh.cells[2].text = "Document Types"
            for c in rh.cells:
                c.paragraphs[0].runs[0].bold = True
            for rule in group_rules:
                row = tbl.add_row()
                row.cells[0].text = rule["rule_id"]
                row.cells[1].text = rule.get("rule_name", "")
                doc_labels = [DOC_TYPE_LABELS.get(d, d) for d in rule.get("doc_scope", [])]
                row.cells[2].text = ", ".join(doc_labels)[:80] or "—"
            doc.add_paragraph()

    # Summary of rule counts
    doc.add_heading("Rule Count Summary", level=1)
    doc.add_paragraph(f"This manual documents {len(rules)} rules implemented in REIMS2. The \"377 Rules Active\" count shown in the UI may include dynamic sub-rules (e.g., per-account or per-escrow-type) that execute at runtime.")
    sum_table = doc.add_table(rows=0, cols=2)
    sum_table.style = "Table Grid"
    sh = sum_table.add_row()
    sh.cells[0].text = "Category"
    sh.cells[1].text = "Count"
    for c in sh.cells:
        c.paragraphs[0].runs[0].bold = True
    for gname in group_order:
        if gname in groups:
            row = sum_table.add_row()
            row.cells[0].text = gname
            row.cells[1].text = str(len(groups[gname]))
    sum_table.add_row().cells[0].text = "TOTAL"
    sum_table.rows[-1].cells[0].paragraphs[0].runs[0].bold = True
    sum_table.rows[-1].cells[1].text = str(len(rules))
    doc.add_paragraph()

    doc.add_heading("Additional Reference", level=1)
    doc.add_paragraph("For implementation details, see RULES_COVERAGE_MATRIX.md and RULES_MAPPING.md in the documentation.")
    cat_table = doc.add_table(rows=0, cols=3)
    cat_table.style = "Table Grid"
    hdr = cat_table.add_row()
    hdr.cells[0].text = "Category"
    hdr.cells[1].text = "Rule IDs"
    hdr.cells[2].text = "Description"
    for c in hdr.cells:
        c.paragraphs[0].runs[0].bold = True
    categories = [
        ("Audit Rules", "AUDIT-1 to AUDIT-55", "Balance sheet equation, cash reconciliation, net income three-way, mortgage, rent roll, working capital, D&A, CapEx, debt service, covenant, variance alerts"),
        ("Data Quality", "DQ-1 to DQ-33", "Field completeness, null checks, accounting equation validation, account consistency"),
        ("Analytics", "ANALYTICS-1 to ANALYTICS-33", "NOI, margin, CoC, cap rate, occupancy, LTV, DSCR, ROA, ROE, efficiency metrics"),
        ("Covenant", "COVENANT-1 to COVENANT-6", "DSCR, LTV, liquidity, occupancy, tenant concentration, reporting deadlines"),
        ("Forensic Anomaly", "FA-1 to FA-7, FA-WC, FA-MORT, FA-RR", "Cash flow consistency, Benford's law, duplicate round numbers, working capital, mortgage, rent roll forensics"),
        ("Rent Roll-BS", "RRBS-1 to RRBS-4", "Security deposits floor, A/R reasonableness, prepaid rent, lease roster completeness"),
    ]
    for cat, ids, desc in categories:
        row = cat_table.add_row()
        row.cells[0].text = cat
        row.cells[1].text = ids
        row.cells[2].text = desc
    doc.add_paragraph()

    # Document types and line items reference
    doc.add_heading("Document Types & Key Line Items Reference", level=1)
    doc.add_paragraph("Rules reference line items from the following document types:")
    ref_table = doc.add_table(rows=0, cols=2)
    ref_table.style = "Table Grid"
    ref_hdr = ref_table.add_row()
    ref_hdr.cells[0].text = "Document Type"
    ref_hdr.cells[1].text = "Key Line Items / Account Codes"
    for c in ref_hdr.cells:
        c.paragraphs[0].runs[0].bold = True
    refs = [
        ("Balance Sheet", "1999-0000 TOTAL ASSETS; 2999-0000 Total Liabilities; 3999-0000 Total Capital; 0499-9000 Total Current Assets; 0122-0000 Cash Operating; 0510-0000 Land; 1230-0000 Accum Depr; 0710-0000 5 Year Improvements; 0950-0000 TI Improvements; 0815-0000 Roof; 0816-0000 HVAC; 1210-0000 Deposits; 1920-0000 Loan Costs"),
        ("Income Statement", "TOTAL INCOME; NET OPERATING INCOME; NET INCOME; Total Expenses; Property Tax; Insurance; Depreciation; Amortization; Mortgage Interest"),
        ("Cash Flow", "BEGINNING CASH; NET CHANGE IN CASH; ENDING CASH; Operating/Investing/Financing totals; Net Income; Depreciation add-back"),
        ("Mortgage Statement", "Principal Balance; Interest; Escrow; Payment components; YTD totals"),
        ("Rent Roll", "Tenant units; Monthly/Annual rent; Occupied/Vacant area; Total monthly rent"),
    ]
    for dt, items in refs:
        row = ref_table.add_row()
        row.cells[0].text = dt
        row.cells[1].text = items
    doc.add_paragraph()

    # Maintenance section
    doc.add_heading("Maintenance Information", level=1)
    doc.add_paragraph("Maintained by: Finance Team")
    doc.add_paragraph("Last Updated: January 31, 2026")
    doc.add_paragraph(f"Document Version: 1.1 — Documents all {len(rules)} rules")

    # Save
    out_dir = Path("/home/hsthind/REIMS - Documents")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "REIMS2_Rules_End_User_Manual.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return str(out_path)


def natural_sort_key(s):
    """Natural sort for rule IDs like BS-1, BS-10, BS-2"""
    parts = re.split(r"(\d+)", s)
    return [int(p) if p.isdigit() else p.lower() for p in parts]


if __name__ == "__main__":
    main()
