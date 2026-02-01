#!/usr/bin/env python3
"""
Generate Market Intelligence User Manual

Creates a comprehensive .docx user manual for the Market Intelligence feature,
including purpose, components, tabs, and end-user workflows.

Output: /home/hsthind/REIMS - Documents/REIMS2_Market_Intelligence_User_Manual.docx
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Use REAL screenshots from this folder. See SCREENSHOT_CAPTURE_GUIDE.md for capture instructions.
SCREENSHOTS_DIR = Path("/home/hsthind/REIMS - Documents/reims-screenshots")

# Expected filenames for Market Intelligence tab screenshots (must be captured from live app)
TAB_SCREENSHOTS = {
    "Demographics": "market-intelligence-demographics.png",
    "Economic Indicators": "market-intelligence-economic.png",
    "Location Intelligence": "market-intelligence-location.png",
    "ESG Assessment": "market-intelligence-esg.png",
    "Forecasts": "market-intelligence-forecasts.png",
    "Competitive Analysis": "market-intelligence-competitive.png",
    "AI Insights": "market-intelligence-ai-insights.png",
    "Data Lineage": "market-intelligence-data-lineage.png",
}


def add_screenshot(doc, filename, caption, width=5.5):
    """Add real screenshot if it exists; otherwise add placeholder with capture instructions."""
    path = SCREENSHOTS_DIR / filename
    if path.exists():
        try:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(path), width=Inches(width))
            if caption:
                cap = doc.add_paragraph()
                cap.add_run(caption).italic = True
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            doc.add_paragraph(f"[Insert screenshot: {filename}]")
    else:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("⚠️ INSERT REAL SCREENSHOT: ").bold = True
        p.add_run(f"Capture from the live REIMS app. Save as '{filename}' in reims-screenshots folder. See SCREENSHOT_CAPTURE_GUIDE.md for exact steps.")
        p.paragraph_format.space_before = Pt(6)
    doc.add_paragraph()


def add_heading_with_decision(doc, title, content, decisions=None):
    """Add a section with optional 'Decisions you can make' callout."""
    doc.add_heading(title, level=2)
    doc.add_paragraph(content)
    if decisions:
        p = doc.add_paragraph()
        p.add_run("Decisions you can make: ").bold = True
        p.add_run(decisions)
        p.paragraph_format.space_before = Pt(6)
    doc.add_paragraph()


def main():
    doc = Document()
    doc.add_heading("REIMS2 Market Intelligence", 0)
    doc.add_paragraph()
    doc.add_heading("User Manual", level=1)
    doc.add_paragraph(
        "This manual explains how to use the Market Intelligence feature to access comprehensive market data, "
        "forecasts, and AI-generated insights for your properties to support acquisition, valuation, and strategic decisions."
    )
    doc.add_paragraph()

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "Purpose of Market Intelligence",
        "How to Access Market Intelligence",
        "Hub Layout Overview",
        "Header & Controls (Back, Refresh, Status)",
        "Tabs: Demographics, Economic Indicators, Location Intelligence, ESG Assessment, Forecasts, Competitive Analysis, AI Insights, Data Lineage",
        "What End Users Should Look At",
        "Recommended Workflow",
    ]
    for item in toc_items:
        doc.add_paragraph(f"• {item}", style="List Bullet")
    doc.add_paragraph()

    # Purpose
    doc.add_heading("Purpose of Market Intelligence", level=1)
    doc.add_paragraph(
        "Market Intelligence provides property-level market data, analytics, and AI insights to support investment decisions. It enables you to:"
    )
    bullets = [
        "Assess demographics: population, median income, home values, rent, unemployment, education, housing mix",
        "Monitor economic indicators: GDP growth, unemployment, inflation, Fed funds rate, mortgage rates, recession probability, MSA metrics",
        "Evaluate location: walkability, transit, bike scores; amenity counts (grocery, restaurants, schools, hospitals, parks); crime index; school ratings",
        "Assess ESG risks: environmental (flood, wildfire, earthquake, climate, emissions), social (crime, schools, inequality, diversity), governance (zoning, permits, tax, legal)",
        "View forecasts: 12-month predictions for rent, occupancy, cap rate, and property value with confidence intervals",
        "Analyze competition: submarket position (rent, occupancy, quality, value percentiles); competitive threats; submarket trends (rent growth, supply, absorption)",
        "Get AI insights: SWOT analysis, investment recommendation (BUY/HOLD/SELL), risk assessment, opportunities, market trend synthesis",
        "Audit data lineage: source, vintage, fetch date, confidence for each data category",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()

    # How to Access
    doc.add_heading("How to Access Market Intelligence", level=1)
    doc.add_paragraph(
        "Market Intelligence is property-specific. Navigate to the Properties page, select a property, and open Market Intelligence "
        "(e.g., click the Market Intelligence button or link). The route is #market-intelligence/{property_code} (e.g., #market-intelligence/ESP001). "
        "You must have a valid property code to view Market Intelligence."
    )
    doc.add_paragraph()

    # Layout Visual
    doc.add_heading("Hub Layout Overview", level=1)
    add_screenshot(doc, "market-intelligence-demographics.png", "Figure 1: Market Intelligence – Demographics tab (or main layout)")
    doc.add_paragraph()

    # Header & Controls
    doc.add_heading("Header & Controls", level=1)
    doc.add_paragraph()
    add_heading_with_decision(
        doc,
        "Back Button",
        "Returns to the previous page (e.g., Properties list).",
        "When to use: To return to property selection or the properties overview.",
    )
    add_heading_with_decision(
        doc,
        "Property Code",
        "Shows the property code for the current Market Intelligence view (e.g., ESP001).",
        "Which property is this data for?",
    )
    add_heading_with_decision(
        doc,
        "Last Updated Chip",
        "Displays when the data was last refreshed. Green = recent; amber = data may be stale (older than 24 hours).",
        "Is my data fresh enough for decision-making? Do I need to refresh?",
    )
    add_heading_with_decision(
        doc,
        "Status Chip",
        "Shows refresh status: success (all categories loaded), partial (some categories failed), failure.",
        "Was the last refresh complete? Should I retry?",
    )
    add_heading_with_decision(
        doc,
        "Refresh All",
        "Fetches or regenerates all market intelligence categories for the property. Use when data is missing or stale. Refresh can take 30–60 seconds.",
        "When to use: When opening for the first time, when data is stale, or after property/location changes.",
    )
    doc.add_paragraph()

    # Tabs
    doc.add_heading("Tabs & Data Categories", level=1)
    doc.add_paragraph()

    tabs = [
        {
            "name": "Demographics",
            "icon": "📍",
            "content": (
                "Census and ACS demographics for the property's geography. Key metrics: Population, Median Household Income, "
                "Median Home Value, Median Gross Rent, Unemployment Rate, Median Age, College Educated %, and Housing Units by type "
                "(single-family, multifamily 2–4, 5–9, 10–19, 20–49, 50+). Geography: State, County, Tract FIPS codes. "
                "Data lineage: source (e.g., Census ACS), vintage, confidence, fetched date."
            ),
            "decisions": "Is this a strong demographic market for multifamily? How does income and rent compare to our targets?",
        },
        {
            "name": "Economic Indicators",
            "icon": "📈",
            "content": (
                "FRED and MSA economic data: GDP Growth, Unemployment Rate, Inflation Rate, Fed Funds Rate, 30-Year Mortgage Rate, "
                "Recession Probability, MSA Unemployment, MSA GDP. Values and dates shown; trends indicated. Helps assess macro and local economic conditions."
            ),
            "decisions": "Are interest rates favorable? Is the local economy strong? What is recession risk?",
        },
        {
            "name": "Location Intelligence",
            "icon": "🗺️",
            "content": (
                "Walk Score, Transit Score, Bike Score (0–100). Amenities: grocery stores, restaurants, schools, hospitals, parks within radius. "
                "Transit access: bus stops, subway/rail, commute time to downtown. Crime index, school rating average. Drive times. "
                "May include map with isochrones (travel-time contours)."
            ),
            "decisions": "Is this a walkable, transit-friendly location? How do amenities and safety compare to targets?",
        },
        {
            "name": "ESG Assessment",
            "icon": "🌱",
            "content": (
                "Environmental: flood risk score/zone, wildfire risk, earthquake risk, climate risk composite, energy efficiency, emissions intensity. "
                "Social: crime score, school quality, income inequality (Gini), diversity index, community health. "
                "Governance: zoning compliance, permit history, tax delinquency risk, legal issues count, regulatory risk. "
                "Composite ESG score and grade (A–D)."
            ),
            "decisions": "What are the main ESG risks? Would lenders or investors flag environmental or governance issues?",
        },
        {
            "name": "Forecasts",
            "icon": "📊",
            "content": (
                "12-month forecasts: Rent, Occupancy, Cap Rate, Property Value. Each forecast shows: predicted value, change %, "
                "95% confidence interval, model used, accuracy/MAE. Supports budgeting, underwriting, and exit planning."
            ),
            "decisions": "What rent growth do we expect? Is occupancy trending up or down? What value change is forecast?",
        },
        {
            "name": "Competitive Analysis",
            "icon": "⚖️",
            "content": (
                "Submarket Position: rent, occupancy, quality, value percentiles (vs. city/portfolio). Competitive Threats: nearby properties, "
                "distance, threat score, advantages/disadvantages. Submarket Trends: rent growth CAGR, occupancy trend, new supply pipeline, "
                "absorption rate, months of supply. May include LLM narrative: positioning summary, differentiation factors, pricing power analysis, strategic recommendations."
            ),
            "decisions": "How do we rank vs. peers? Which competitors are threats? What is the supply/demand outlook?",
        },
        {
            "name": "AI Insights",
            "icon": "🤖",
            "content": (
                "SWOT Analysis: strengths, weaknesses, opportunities, threats. Investment Recommendation: BUY, HOLD, or SELL with confidence score, rationale, key factors. "
                "Risk assessment, opportunities list, market trend synthesis. Data coverage: which categories are present vs. missing, affecting insight quality."
            ),
            "decisions": "What is the AI recommendation? What are the main risks and opportunities? Is data coverage sufficient?",
        },
        {
            "name": "Data Lineage",
            "icon": "📜",
            "content": (
                "Audit trail for market intelligence data. Table of records: source, category, vintage, fetched date, status (success/partial/failure), "
                "confidence, records fetched, errors. Filter by category. Use to verify data freshness and reliability."
            ),
            "decisions": "When was each category last fetched? Which sources failed? Is the data trustworthy?",
        },
    ]

    for t in tabs:
        add_heading_with_decision(doc, f"{t['icon']} {t['name']} Tab", t["content"], t["decisions"])
        # Add tab screenshot if available (real screenshot from reims-screenshots folder)
        filename = TAB_SCREENSHOTS.get(t["name"])
        if filename:
            add_screenshot(doc, filename, f"Figure: {t['name']} tab", width=6.0)

    # Fetch buttons for missing data
    doc.add_heading("Fetching Missing Data", level=1)
    doc.add_paragraph(
        "If a tab shows no data, a 'Fetch [Category] Data' or 'Generate [Category]' button appears. "
        "Click it to request that category. Some categories (e.g., Location, ESG, Forecasts, Competitive, AI Insights) may auto-fetch on load "
        "or require a manual refresh. Refresh All updates all categories."
    )
    doc.add_paragraph()

    # What to Look At
    doc.add_heading("What End Users Should Look At", level=1)
    doc.add_paragraph("When evaluating a property for acquisition, underwriting, or reporting, focus on:")
    checklist = [
        "Demographics – median income, rent, and housing mix; does the market support our rent targets?",
        "Economic Indicators – mortgage rates, recession probability; is the macro environment favorable?",
        "Location Intelligence – walk/transit scores, amenities, crime; does the location meet our criteria?",
        "ESG Assessment – flood, climate, governance risks; any deal-breakers for lenders or investors?",
        "Forecasts – rent and value outlook; are projections aligned with our underwriting?",
        "Competitive Analysis – submarket position and threats; how do we stack up vs. peers?",
        "AI Insights – BUY/HOLD/SELL recommendation and SWOT; what are the key risks and opportunities?",
        "Data Lineage – freshness and status; is the underlying data reliable?",
    ]
    for c in checklist:
        doc.add_paragraph(c, style="List Bullet")
    doc.add_paragraph()

    # Workflow
    doc.add_heading("Recommended Workflow", level=1)
    doc.add_paragraph("1. Navigate to Properties and select a property")
    doc.add_paragraph("2. Open Market Intelligence for that property (#market-intelligence/{property_code})")
    doc.add_paragraph("3. If data is missing or stale, click Refresh All")
    doc.add_paragraph("4. Review Demographics and Economic Indicators for market context")
    doc.add_paragraph("5. Check Location Intelligence and ESG for risk factors")
    doc.add_paragraph("6. Review Forecasts for rent, occupancy, and value outlook")
    doc.add_paragraph("7. Use Competitive Analysis to understand positioning")
    doc.add_paragraph("8. Read AI Insights for investment recommendation and SWOT")
    doc.add_paragraph("9. Use Data Lineage to verify data quality and freshness")
    doc.add_paragraph()

    # Maintenance
    doc.add_heading("Maintenance", level=1)
    doc.add_paragraph("Last Updated: January 31, 2026")
    doc.add_paragraph("Document Version: 1.0")

    out_dir = Path("/home/hsthind/REIMS - Documents")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "REIMS2_Market_Intelligence_User_Manual.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return str(out_path)


if __name__ == "__main__":
    main()
