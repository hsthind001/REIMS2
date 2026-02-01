#!/usr/bin/env python3
"""
Generate Forensic Audit Dashboard User Manual

Creates a comprehensive .docx user manual for the Forensic Audit Dashboard,
including purpose, components, tabs, and end-user workflows.

Output: /home/hsthind/REIMS - Documents/REIMS2_Forensic_Audit_Dashboard_User_Manual.docx
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Use REAL screenshots from this folder. See SCREENSHOT_CAPTURE_GUIDE.md for capture instructions.
SCREENSHOTS_DIR = Path("/home/hsthind/REIMS - Documents/reims-screenshots")


def add_screenshot(doc, filename, caption, width=5.5):
    """Add real screenshot if it exists; otherwise add placeholder with capture instructions."""
    path = SCREENSHOTS_DIR / filename
    if path.exists():
        try:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(path), width=Inches(width))
            if caption:
                cap = doc.add_paragraph()
                cap.add_run(caption).italic = True
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            doc.add_paragraph(f"[Insert screenshot: {filename}]")
    else:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("⚠️ INSERT REAL SCREENSHOT: ").bold = True
        p.add_run(f"Capture from the live REIMS app. Save as '{filename}' in reims-screenshots folder. See SCREENSHOT_CAPTURE_GUIDE.md for exact steps.")
        p.paragraph_format.space_before = Pt(6)
    doc.add_paragraph()


def add_heading_with_decision(doc, title, content, decisions=None):
    """Add a section with optional 'Decisions you can make' callout."""
    doc.add_heading(title, level=2)
    doc.add_paragraph(content)
    if decisions:
        p = doc.add_paragraph()
        p.add_run("Decisions you can make: ").bold = True
        p.add_run(decisions)
        p.paragraph_format.space_before = Pt(6)
    doc.add_paragraph()


def main():
    doc = Document()
    doc.add_heading("REIMS2 Forensic Audit Dashboard", 0)
    doc.add_paragraph()
    doc.add_heading("User Manual", level=1)
    doc.add_paragraph(
        "This manual explains how to use the Forensic Audit Dashboard to perform Big 5 accounting firm-level "
        "comprehensive analysis of your property's financial data, identify risks, and make informed decisions."
    )
    doc.add_paragraph()

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "Purpose of the Forensic Audit Dashboard",
        "How to Access the Forensic Audit Dashboard",
        "Hub Layout Overview",
        "Header & Controls (Property, Period, Refresh, Export, Run Audit)",
        "Top Summary Cards (Overall Health, Audit Opinion, Overall Status)",
        "Key Financial Metrics (DSCR, LTV, Fraud Risk, Reconciliation Pass Rate)",
        "Traffic Light Metrics & Audit Drilldowns",
        "Priority Risks & Action Items",
        "Tabs: Overview, Math Integrity, Performance, Fraud Detection, Covenants, Tenant Risk, Collections, Documents, Reconciliation, History",
        "What End Users Should Look At",
        "Recommended Workflow",
    ]
    for item in toc_items:
        doc.add_paragraph(f"• {item}", style="List Bullet")
    doc.add_paragraph()

    # Purpose
    doc.add_heading("Purpose of the Forensic Audit Dashboard", level=1)
    doc.add_paragraph(
        "The Forensic Audit Dashboard provides executive-level, Big 5 accounting firm-style comprehensive analysis "
        "of property financial data. It enables you to:"
    )
    bullets = [
        "Obtain an overall health score and audit opinion (Clean, Qualified, Adverse, Disclaimer)",
        "Monitor key covenants: DSCR, LTV, Interest Coverage, Liquidity",
        "Detect fraud indicators: Benford's Law deviations, round numbers, duplicate payments, cash conversion anomalies",
        "Assess tenant risk: concentration, lease rollover, occupancy, credit quality",
        "Evaluate collections quality: DSO, revenue quality, A/R aging",
        "Verify document completeness: required audit inputs for the property/period",
        "Check math integrity: internal calculation consistency across Balance Sheet, Income Statement, Cash Flow",
        "Benchmark performance: NOI margin, OpEx ratio, CapEx ratio, same-store growth",
        "Review cross-document reconciliation: 9+ tie-outs between Balance Sheet, Income Statement, Rent Roll, Mortgage Statement, etc.",
        "Track audit history: health score and covenant trends over time",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()

    # How to Access
    doc.add_heading("How to Access the Forensic Audit Dashboard", level=1)
    doc.add_paragraph(
        "Navigate to Operations and open the Forensic Audit Dashboard (route: #forensic-audit-dashboard). "
        "You can also access it from Quality Control, Risk Intelligence, Audit History, Reconciliation Results, "
        "Document Completeness, Collections Quality, Tenant Risk, Fraud Detection, Covenant Compliance, "
        "or Math Integrity dashboards via the 'Forensic Audit' or 'Back to Overview' link."
    )
    doc.add_paragraph()

    # Layout Visual
    doc.add_heading("Hub Layout Overview", level=1)
    add_screenshot(doc, "forensic-audit-overview.png", "Figure 1: Forensic Audit Dashboard Overview")
    doc.add_paragraph()

    # Header & Controls
    doc.add_heading("Header & Controls", level=1)
    doc.add_paragraph()
    add_heading_with_decision(
        doc,
        "Property Selector",
        "Select the property to audit. Results are scoped to the selected property. Selection is persisted for your next visit.",
        "Which property should I audit?",
    )
    add_heading_with_decision(
        doc,
        "Period Selector",
        "Select the financial period (e.g., 2025-01). Periods load based on the selected property. The most recent complete period is often auto-selected.",
        "Which reporting period should I analyze?",
    )
    add_heading_with_decision(
        doc,
        "Refresh",
        "Reload the scorecard and all detail panels without re-running the full audit. Use after data changes to refresh displayed metrics.",
        "When to use: After uploading documents or updating extractions.",
    )
    add_heading_with_decision(
        doc,
        "Export PDF / Export Excel",
        "Download a forensic audit report in PDF or Excel format for the selected property and period. Useful for board packs, lender reporting, or audit files.",
        "When to use: For lender reporting, audit documentation, or executive summaries.",
    )
    add_heading_with_decision(
        doc,
        "Audit History",
        "Navigate to the Audit History tab to view health score and covenant trends over time for the selected property.",
        "When to use: To review trends and identify deteriorating metrics.",
    )
    add_heading_with_decision(
        doc,
        "Run Audit",
        "Execute the full forensic audit for the selected property and period. The audit runs asynchronously; progress is shown. On completion, the scorecard and all drilldown data refresh automatically.",
        "When to run: When no scorecard exists yet, after major data changes, or before month-end/quarter-end reporting.",
    )
    doc.add_paragraph()

    # Top Summary Cards
    doc.add_heading("Top Summary Cards", level=1)
    doc.add_paragraph()
    add_heading_with_decision(
        doc,
        "Overall Health Score",
        "A 0–100% gauge representing the aggregate health of the property's financial data across all audit dimensions. Higher is better. Generated timestamp is shown.",
        "Is the property in good financial health? Is it improving or deteriorating?",
    )
    add_heading_with_decision(
        doc,
        "Audit Opinion",
        "Simulated audit opinion: Clean, Qualified, Adverse, or Disclaimer. Reflects the severity of issues detected. Optional auditor notes may appear.",
        "What would an external auditor conclude? Do I need to disclose qualifications?",
    )
    add_heading_with_decision(
        doc,
        "Overall Status (Traffic Light)",
        "Pass (green), Warning (yellow), or Fail (red). Shows counts of Passed, Warning, and Failed reconciliations for quick at-a-glance status.",
        "Can I proceed with closing, or must I fix issues first?",
    )
    doc.add_paragraph()

    # Key Financial Metrics
    doc.add_heading("Key Financial Metrics", level=1)
    doc.add_paragraph()
    add_heading_with_decision(
        doc,
        "DSCR (Debt Service Coverage Ratio)",
        "Current DSCR value vs. covenant threshold. Shows status (Pass/Warning/Fail) and covenant target.",
        "Are we in compliance with the lender's DSCR covenant?",
    )
    add_heading_with_decision(
        doc,
        "LTV Ratio",
        "Loan-to-Value ratio vs. covenant. Status indicates compliance.",
        "Are we within the LTV covenant limit?",
    )
    add_heading_with_decision(
        doc,
        "Fraud Risk",
        "Overall fraud risk level (Low, Medium, High) based on Benford's Law, round numbers, duplicates, and cash conversion.",
        "Should I investigate further for potential fraud or manipulation?",
    )
    add_heading_with_decision(
        doc,
        "Reconciliation Pass Rate",
        "Percentage of cross-document reconciliations that passed. Shows passed/total count and status.",
        "How consistent is our data across documents? Where are the tie-out failures?",
    )
    doc.add_paragraph()

    # Traffic Light Metrics & Audit Drilldowns
    doc.add_heading("Traffic Light Metrics & Audit Drilldowns", level=1)
    doc.add_paragraph(
        "Traffic Light Metrics: Individual metrics (e.g., DSCR, LTV, occupancy) with current value, target, trend indicator, and status."
    )
    doc.add_paragraph(
        "Audit Drilldowns: Cards for Reconciliations, Fraud Detection, Covenant Compliance, Tenant Risk, Collections Quality, "
        "Math Integrity, Document Completeness, and Performance Benchmarking. Each card shows a summary and a 'View' button "
        "to open the corresponding detailed tab."
    )
    p = doc.add_paragraph()
    p.add_run("Decisions you can make: ").bold = True
    p.add_run("Which dimension needs the most attention? Where should I drill down first?")
    doc.add_paragraph()

    # Priority Risks & Action Items
    doc.add_heading("Priority Risks & Action Items", level=1)
    doc.add_paragraph(
        "Priority Risks: List of risks ranked by severity with category, description, action required, owner, due date, and financial impact. "
        "Use this to prioritize remediation work."
    )
    doc.add_paragraph(
        "Action Items: Tasks assigned to owners with priority (URGENT, HIGH, MEDIUM, LOW), due date, and status (COMPLETED, IN_PROGRESS, PENDING). "
        "Track follow-up and ensure nothing falls through the cracks."
    )
    p = doc.add_paragraph()
    p.add_run("Decisions you can make: ").bold = True
    p.add_run("What should I do first? Who owns each action? Are we on track for due dates?")
    doc.add_paragraph()

    # Tabs
    doc.add_heading("Tabs & Specialized Dashboards", level=1)
    doc.add_paragraph()

    tabs = [
        {
            "name": "Overview",
            "icon": "📊",
            "content": (
                "Executive summary of the Forensic Audit Dashboard. Shows Overall Health Score, Audit Opinion, Overall Status, "
                "Key Financial Metrics (DSCR, LTV, Fraud Risk, Reconciliation Pass Rate), Traffic Light Metrics, Audit Drilldowns "
                "(Reconciliations, Fraud Detection, Covenant Compliance, Tenant Risk, Collections, Math Integrity, Document Completeness, Performance), "
                "Priority Risks, Action Items, and Financial Summary (Total Revenue, Net Income, NOI, Cash Balance)."
            ),
            "decisions": "What is the overall picture? What needs immediate attention?",
        },
        {
            "name": "Math Integrity",
            "icon": "🔢",
            "content": (
                "Internal calculation checks for Balance Sheet, Income Statement, and Cash Flow. Validates that subtotals, totals, "
                "and key equations (e.g., Assets = Liabilities + Equity) are internally consistent. Shows passed/total checks and status."
            ),
            "decisions": "Are there arithmetic or formula errors in our core financial documents?",
        },
        {
            "name": "Performance Benchmarking",
            "icon": "📈",
            "content": (
                "Benchmarks: Same-store growth, NOI margin, Operating Expense ratio, CapEx ratio. Compares current period to targets "
                "and shows trend. Helps assess operational efficiency and investment performance."
            ),
            "decisions": "How is the property performing vs. targets? Are OpEx or CapEx out of line?",
        },
        {
            "name": "Fraud Detection",
            "icon": "🚨",
            "content": (
                "Benford's Law analysis, round number detection, duplicate payment flags, and cash conversion ratio. Identifies anomalies "
                "that may indicate manipulation or fraud. Shows overall risk level and detailed test results."
            ),
            "decisions": "Should I escalate for fraud investigation? Which transactions need manual review?",
        },
        {
            "name": "Covenant Compliance",
            "icon": "📋",
            "content": (
                "DSCR, LTV, Interest Coverage Ratio, and Liquidity ratios with covenant thresholds and breach status. "
                "Shows calculated values, targets, and Pass/Fail for each covenant."
            ),
            "decisions": "Are we in breach? Which covenants are at risk? Do I need to notify the lender?",
        },
        {
            "name": "Tenant Risk",
            "icon": "🏢",
            "content": (
                "Tenant concentration (Top 5 tenants %), lease rollover (12-month rollover %), occupancy, and credit quality. "
                "Identifies concentration risk and rollover exposure."
            ),
            "decisions": "Is tenant concentration too high? What is our rollover exposure? Should I diversify?",
        },
        {
            "name": "Collections Quality",
            "icon": "💰",
            "content": (
                "Revenue quality score, DSO (Days Sales Outstanding), cash conversion, and A/R aging. "
                "Evaluates how quickly revenue converts to cash and aging of receivables."
            ),
            "decisions": "Is collections healthy? Are we collecting too slowly? Is A/R aging a concern?",
        },
        {
            "name": "Document Completeness",
            "icon": "📄",
            "content": (
                "Required audit document availability. Shows completeness percentage and lists missing documents. "
                "Ensures all necessary inputs (Balance Sheet, Income Statement, Cash Flow, Rent Roll, Mortgage Statement, etc.) are present."
            ),
            "decisions": "What documents are missing? Can I run a full audit with current data?",
        },
        {
            "name": "Reconciliation Results",
            "icon": "🔄",
            "content": (
                "All cross-document reconciliation tie-outs (9+ combinations). Shows pass/fail for each pair (e.g., Balance Sheet ↔ Income Statement, "
                "Rent Roll ↔ Balance Sheet). Expand each to see details of matched items and variances."
            ),
            "decisions": "Which document pairs have variances? Where do I need to investigate discrepancies?",
        },
        {
            "name": "Audit History",
            "icon": "📜",
            "content": (
                "Trend view of audit results over time. Shows Overall Health Score trend, DSCR trend, Occupancy trend, and other metrics across periods. "
                "No period selector—property-level history. Use to spot deteriorating or improving trends."
            ),
            "decisions": "Is the property improving or declining over time? What is the trend in DSCR and occupancy?",
        },
    ]

    for t in tabs:
        add_heading_with_decision(doc, f"{t['icon']} {t['name']} Tab", t["content"], t["decisions"])

    # What to Look At
    doc.add_heading("What End Users Should Look At", level=1)
    doc.add_paragraph("Before month-end close, board reporting, or lender reviews, focus on:")
    checklist = [
        "Overall Health Score – ensure it is in an acceptable range for your risk appetite",
        "Audit Opinion – address any Qualified or Adverse opinion before external reporting",
        "Overall Status – resolve Fail (red) before closing",
        "DSCR and LTV – confirm covenant compliance; escalate breaches immediately",
        "Fraud Risk – investigate High risk; review flagged transactions",
        "Priority Risks – address High and Critical items; assign owners and due dates",
        "Action Items – ensure URGENT and HIGH items are completed or in progress",
        "Document Completeness – obtain missing documents before running final audit",
        "Reconciliation Results – resolve material variances between documents",
        "Audit History – spot trends and address deteriorating metrics",
    ]
    for c in checklist:
        doc.add_paragraph(c, style="List Bullet")
    doc.add_paragraph()

    # Workflow
    doc.add_heading("Recommended Workflow", level=1)
    doc.add_paragraph("1. Select Property and Period")
    doc.add_paragraph("2. Verify Document Completeness (Documents tab) – obtain missing documents if needed")
    doc.add_paragraph("3. Click Run Audit and wait for completion (progress shown)")
    doc.add_paragraph("4. Review Overview – Overall Health Score, Audit Opinion, Key Metrics")
    doc.add_paragraph("5. Check Priority Risks and Action Items – triage and assign")
    doc.add_paragraph("6. Drill into failing areas – Covenant Compliance, Reconciliation Results, Fraud Detection, etc.")
    doc.add_paragraph("7. Resolve critical issues and Refresh to update scorecard")
    doc.add_paragraph("8. Export PDF or Excel for reporting/audit files")
    doc.add_paragraph("9. Use Audit History to monitor trends over time")
    doc.add_paragraph()

    # Maintenance
    doc.add_heading("Maintenance", level=1)
    doc.add_paragraph("Last Updated: January 31, 2026")
    doc.add_paragraph("Document Version: 1.0")

    out_dir = Path("/home/hsthind/REIMS - Documents")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "REIMS2_Forensic_Audit_Dashboard_User_Manual.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return str(out_path)


if __name__ == "__main__":
    main()
