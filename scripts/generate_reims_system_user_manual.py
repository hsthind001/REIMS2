#!/usr/bin/env python3
"""
Generate REIMS 2.0 Complete System User Manual

Creates a comprehensive .docx user manual for the entire REIMS system.
Uses Option A: Placeholder instructions for screenshots. User adds real screenshots manually.

Output: /home/hsthind/REIMS - Documents/REIMS2_Complete_System_User_Manual.docx
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_screenshot_placeholder(doc, caption, route_or_instruction):
    """Add a clear placeholder for the user to insert a real screenshot."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("📷 SCREENSHOT PLACEHOLDER").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"[Insert screenshot: {caption}]",
        style="Intense Quote"
    )
    inst = doc.add_paragraph()
    inst.add_run("How to capture: ").bold = True
    inst.add_run(route_or_instruction)
    inst.paragraph_format.space_after = Pt(12)
    doc.add_paragraph()


def add_heading_with_decision(doc, title, content, decisions=None):
    """Add a section with optional 'Decisions you can make' callout."""
    doc.add_heading(title, level=2)
    doc.add_paragraph(content)
    if decisions:
        p = doc.add_paragraph()
        p.add_run("Decisions you can make: ").bold = True
        p.add_run(decisions)
        p.paragraph_format.space_before = Pt(6)


def main():
    doc = Document()
    doc.add_heading("REIMS 2.0", 0)
    doc.add_heading("Complete System User Manual", level=1)
    doc.add_paragraph()
    doc.add_paragraph(
        "This manual provides end-to-end guidance for using the REIMS (Real Estate Investment Management System) 2.0 platform. "
        "It covers all main pages, tabs, components, and sub-pages. Use the navigation sidebar and hash routes to reach each area."
    )
    doc.add_paragraph()

    # Screenshot instructions
    doc.add_heading("Adding Screenshots to This Manual", level=1)
    doc.add_paragraph(
        "This document includes placeholder instructions for screenshots. To add real screenshots:"
    )
    doc.add_paragraph("1. Open REIMS in your browser and log in.", style="List Number")
    doc.add_paragraph("2. Navigate to the page or view indicated in each placeholder.", style="List Number")
    doc.add_paragraph("3. Take a screenshot (e.g., Ctrl+Shift+S, or your OS screenshot tool).", style="List Number")
    doc.add_paragraph("4. In Word, delete the placeholder text and insert your screenshot (Insert → Pictures).", style="List Number")
    doc.add_paragraph("5. Resize as needed and add a caption if desired.", style="List Number")
    doc.add_paragraph()

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "Purpose of REIMS",
        "Main Navigation & Sidebar",
        "Insights Hub",
        "Properties",
        "Financials (Reports)",
        "Quality Control (Operations)",
        "Administration",
        "Risk Intelligence",
        "AI Assistant (Natural Language Query)",
        "Hash Routes Quick Reference",
    ]
    for item in toc:
        doc.add_paragraph(f"• {item}", style="List Bullet")
    doc.add_paragraph()

    # Purpose
    doc.add_heading("Purpose of REIMS", level=1)
    doc.add_paragraph(
        "REIMS 2.0 is a Real Estate Investment Management System that helps you manage property portfolios, "
        "financial data, documents, reconciliations, risk, and compliance. Key purposes include:"
    )
    for b in [
        "Portfolio overview: View aggregate health, NOI, value, DSCR, occupancy across properties",
        "Property management: Add properties, view details, market intelligence, tenants, documents",
        "Financial reporting: Statements, variance analysis, chart of accounts, reconciliation",
        "Quality & operations: Document extraction, validation rules, review queue, bulk import",
        "Risk management: Anomalies, alerts, workflow locks, covenant compliance, forensic audit",
        "Administration: Users, roles, organization, billing, audit log, settings",
        "AI: Natural language queries about financial data, formulas, and temporal queries",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()

    # Main Navigation
    doc.add_heading("Main Navigation & Sidebar", level=1)
    add_screenshot_placeholder(
        doc,
        "Main navigation sidebar with all menu items visible",
        "Ensure sidebar is open. Click the sidebar toggle if collapsed. Capture the full left sidebar showing: Insights Hub, Properties, Financials, Quality Control, Administration, Risk Intelligence, AI Assistant."
    )
    doc.add_paragraph("The sidebar provides access to seven main areas:")
    nav_items = [
        ("Insights Hub", "Dashboard, portfolio health, critical alerts, property performance"),
        ("Properties", "Property list, add property, property details, market intelligence"),
        ("Financials", "Statements, variance, chart of accounts, reconciliation, AI assistant"),
        ("Quality Control", "Quality score, tasks, validation rules, import, review queue, documents"),
        ("Administration", "Users, roles, organization, audit, billing, settings"),
        ("Risk Intelligence", "Unified risk workbench, anomalies, alerts, locks, analytics"),
        ("AI Assistant", "Natural language query (NLQ) interface"),
    ]
    for name, desc in nav_items:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)
    doc.add_paragraph("Shortcuts: Ctrl/Cmd+1 (Insights), +2 (Properties), +3 (Financials), +4 (Quality), +5 (Admin), +6 (Risk), +7 (AI).")
    doc.add_paragraph()

    # Insights Hub
    doc.add_heading("Insights Hub", level=1)
    add_screenshot_placeholder(
        doc,
        "Insights Hub – main dashboard",
        "Navigate to Insights Hub (first item in sidebar, or Ctrl/Cmd+1). Ensure no hash in URL. Capture the full dashboard view."
    )
    doc.add_heading("Purpose", level=2)
    doc.add_paragraph("Executive dashboard showing portfolio health, critical alerts, property performance, AI insights, and document status.")
    doc.add_heading("Key Components", level=2)
    comps = [
        ("Portfolio Health", "Overall score, total value, NOI, occupancy, DSCR. Status: excellent/good/fair/poor."),
        ("Critical Alerts", "Alerts requiring attention. Links to property and reports."),
        ("Property Performance", "Per-property value, NOI, DSCR, LTV, occupancy, trends."),
        ("AI Insights", "Risk, opportunity, market, operational insights with confidence."),
        ("Document Matrix", "Document completeness by property and type."),
        ("Quick Actions", "Upload documents, add property, run analysis."),
    ]
    for name, desc in comps:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)
    add_heading_with_decision(doc, "Decisions", "Which properties need attention? What is portfolio health? Where are critical gaps?", None)
    doc.add_paragraph()

    # Properties
    doc.add_heading("Properties", level=1)
    add_screenshot_placeholder(
        doc,
        "Properties – property list",
        "Click Properties in sidebar (or Ctrl/Cmd+2). Hash: (none) or #. Capture the property list/cards."
    )
    doc.add_heading("Property List View", level=2)
    doc.add_paragraph("Lists all properties with summary info. Add Property button opens #add-property.")
    doc.add_heading("Property Detail Tabs", level=2)
    doc.add_paragraph("When a property is selected, tabs appear:")
    for t in ["Overview", "Financials", "Market", "Tenants", "Documents"]:
        doc.add_paragraph(f"  • {t}", style="List Bullet")
    doc.add_heading("Market Intelligence", level=2)
    doc.add_paragraph("Open Market Intelligence for a property: click the link/button or go to #market-intelligence/{property_code} (e.g., #market-intelligence/ESP001).")
    add_screenshot_placeholder(
        doc,
        "Property detail – Overview tab",
        "Select a property. Ensure Overview tab is active. Capture the detail view."
    )
    add_screenshot_placeholder(
        doc,
        "Market Intelligence – Demographics tab",
        "Navigate to #market-intelligence/ESP001 (replace ESP001 with your property code). Capture the Demographics tab content."
    )
    add_heading_with_decision(doc, "Decisions", "Which property to analyze? Is market intelligence favorable? What documents are missing?", None)
    doc.add_paragraph()

    # Financials
    doc.add_heading("Financials (Reports)", level=1)
    add_screenshot_placeholder(
        doc,
        "Financials – main view with tabs",
        "Click Financials in sidebar (or Ctrl/Cmd+3). Hash: #statements, #variance, #chart-of-accounts, #reconciliation, or # (AI tab). Capture the tab bar and main content."
    )
    doc.add_heading("Tabs", level=2)
    fin_tabs = [
        ("AI", "Natural language query. Ask questions about financial data. Default tab."),
        ("Statements", "Balance sheet, income statement, cash flow by property and period."),
        ("Variance", "Variance analysis – actual vs budget/forecast."),
        ("Exit", "Exit strategy scenarios and analysis."),
        ("Chart of Accounts", "Chart of accounts view and management."),
        ("Reconciliation", "Reconciliation sessions and status."),
    ]
    for name, desc in fin_tabs:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)
    doc.add_paragraph("Hash routes: #variance, #statements, #chart-of-accounts, #reconciliation.")
    doc.add_paragraph("Full financial data view: #financial-data?property=X&period=Y")
    add_screenshot_placeholder(
        doc,
        "Financials – Statements tab",
        "Navigate to Financials, click Statements tab. Hash: #statements. Select a property and period. Capture the statements view."
    )
    add_screenshot_placeholder(
        doc,
        "Financials – Variance tab",
        "Navigate to Financials, click Variance tab. Hash: #variance. Capture the variance analysis view."
    )
    add_heading_with_decision(doc, "Decisions", "How is the property performing? What variances exist? Are statements complete?", None)
    doc.add_paragraph()

    # Quality Control
    doc.add_heading("Quality Control (Operations)", level=1)
    add_screenshot_placeholder(
        doc,
        "Quality Control – Quality tab",
        "Click Quality Control in sidebar (or Ctrl/Cmd+4). Hash: (none). Capture the Quality tab with quality score cards."
    )
    doc.add_heading("Tabs", level=2)
    qc_tabs = [
        ("Quality", "Quality score, extraction accuracy, validation pass rate, completeness. Links to Forensic Audit and Review Queue."),
        ("Tasks", "Task dashboard, worker status, task filters, performance metrics, scheduler."),
        ("Validation Rules", "Rule statistics, rule management, configure rules."),
        ("Import", "Bulk import. Link to #bulk-import."),
        ("Review", "Review queue. Link to #review-queue."),
        ("Documents", "Document management and status."),
    ]
    for name, desc in qc_tabs:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)
    doc.add_heading("Financial Integrity Hub", level=2)
    doc.add_paragraph("Navigate to #forensic-reconciliation for live reconciliation, rule validation, and integrity metrics.")
    add_screenshot_placeholder(
        doc,
        "Financial Integrity Hub – Overview tab",
        "Navigate to #forensic-reconciliation. Select property and period. Run reconciliation if needed. Capture the Overview tab with Integrity Score, metrics, and drilldowns."
    )
    doc.add_paragraph("Financial Integrity Hub tabs: Overview, By Document, By Rule, Exceptions, Insights.")
    add_heading_with_decision(doc, "Decisions", "What is data quality? Which rules are failing? What needs review?", None)
    doc.add_paragraph()

    # Bulk Import & Review Queue
    doc.add_heading("Bulk Import", level=1)
    add_screenshot_placeholder(
        doc,
        "Bulk Import page",
        "Navigate to #bulk-import. Capture the upload area and options."
    )
    doc.add_paragraph("Upload documents in bulk. Supports CSV, Excel, and document files.")
    doc.add_paragraph()

    doc.add_heading("Review Queue", level=1)
    add_screenshot_placeholder(
        doc,
        "Review Queue page",
        "Navigate to #review-queue. Optionally add ?severity=critical or ?severity=warning. Capture the queue list."
    )
    doc.add_paragraph("Items requiring manual review. Filter by severity.")
    doc.add_paragraph()

    # Administration
    doc.add_heading("Administration", level=1)
    add_screenshot_placeholder(
        doc,
        "Administration – Users tab",
        "Click Administration in sidebar (or Ctrl/Cmd+5). Hash: (none). Capture the Users tab."
    )
    doc.add_heading("Tabs", level=2)
    admin_tabs = [
        ("Users", "User management, add/edit users, roles."),
        ("Roles", "Role management, permissions."),
        ("Organization", "Organization members, structure."),
        ("Audit", "Audit log of system actions."),
        ("Billing", "Billing overview, history, plans, tenant plans."),
        ("Settings", "System settings and configuration."),
    ]
    for name, desc in admin_tabs:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)
    add_heading_with_decision(doc, "Decisions", "Who has access? What changed? What is our billing status?", None)
    doc.add_paragraph()

    # Risk Intelligence
    doc.add_heading("Risk Intelligence", level=1)
    add_screenshot_placeholder(
        doc,
        "Risk Intelligence – Unified view",
        "Click Risk Intelligence in sidebar (or Ctrl/Cmd+6). Hash: (none). Capture the unified risk workbench."
    )
    doc.add_heading("View Modes", level=2)
    risk_views = [
        ("Unified", "Combined anomalies, alerts, and locks in one table."),
        ("Anomalies", "Statistical anomalies from extraction/validation."),
        ("Alerts", "Alert rules and triggered alerts."),
        ("Locks", "Workflow locks (e.g., document in use)."),
        ("Analytics", "Risk analytics and charts."),
        ("Value Setup", "Value setup for risk calculations."),
    ]
    for name, desc in risk_views:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)
    doc.add_heading("Sub-Pages", level=2)
    doc.add_paragraph("• #anomaly-dashboard – Anomaly dashboard with batch, patterns, all/uncertain tabs")
    doc.add_paragraph("• #anomaly-details?anomaly_id=X – Anomaly detail page")
    doc.add_paragraph("• #alert-rules – Alert rules configuration")
    doc.add_paragraph("• #workflow-locks – Workflow locks view")
    doc.add_paragraph("• #covenant-compliance – Covenant compliance dashboard")
    doc.add_paragraph("• #forensic-audit-dashboard – Forensic Audit Framework (Big 5 style)")
    doc.add_paragraph()
    doc.add_heading("Forensic Audit Dashboard", level=2)
    doc.add_paragraph("Navigate to #forensic-audit-dashboard. Tabs: Overview, Math Integrity, Performance, Fraud Detection, Covenants, Tenant Risk, Collections, Documents, Reconciliation, History.")
    add_screenshot_placeholder(
        doc,
        "Forensic Audit Dashboard – Overview",
        "Navigate to #forensic-audit-dashboard. Select property and period. Run audit if needed. Capture the Overview with Health Score, Audit Opinion, Key Metrics."
    )
    add_heading_with_decision(doc, "Decisions", "What risks need mitigation? Which anomalies to investigate? Are covenants in compliance?", None)
    doc.add_paragraph()

    # AI Assistant
    doc.add_heading("AI Assistant (Natural Language Query)", level=1)
    add_screenshot_placeholder(
        doc,
        "Natural Language Query – main page",
        "Click AI Assistant in sidebar (or Ctrl/Cmd+7). Hash: #nlq-search. Capture the NLQ page with search input and example queries."
    )
    doc.add_paragraph("Ask questions in plain English about financial data, formulas, temporal periods. Example: 'What was cash position in November 2025?'")
    add_heading_with_decision(doc, "Decisions", "What was NOI? How is DSCR calculated? Show revenue for Q4 2025?", None)
    doc.add_paragraph()

    # Rule Configuration
    doc.add_heading("Rule Configuration & Edit", level=1)
    doc.add_paragraph("• #rule-configuration/{rule_id} – Configure a specific rule")
    doc.add_paragraph("• #rule-edit/{rule_id} – Edit rule logic")
    doc.add_paragraph("• #syntax-guide – Syntax guide for rule expressions")
    doc.add_paragraph("Accessed from Financial Integrity Hub (By Rule tab) or Quality Control (Validation Rules).")
    doc.add_paragraph()

    # Hash Routes Reference
    doc.add_heading("Hash Routes Quick Reference", level=1)
    routes_table = doc.add_table(rows=1, cols=2)
    routes_table.style = "Table Grid"
    h = routes_table.rows[0].cells
    h[0].text = "Route"
    h[1].text = "Page / View"
    route_list = [
        ("(none) or #", "Insights Hub / default"),
        ("#bulk-import", "Bulk Import"),
        ("#review-queue", "Review Queue"),
        ("#forensic-reconciliation", "Financial Integrity Hub"),
        ("#market-intelligence/{code}", "Market Intelligence (e.g. ESP001)"),
        ("#variance", "Financials – Variance"),
        ("#statements", "Financials – Statements"),
        ("#chart-of-accounts", "Financials – Chart of Accounts"),
        ("#reconciliation", "Financials – Reconciliation"),
        ("#financial-data?property=&period=", "Full Financial Data view"),
        ("#add-property", "Add Property"),
        ("#anomaly-dashboard", "Anomaly Dashboard"),
        ("#anomaly-details?anomaly_id=", "Anomaly Detail"),
        ("#alert-rules", "Alert Rules"),
        ("#workflow-locks", "Workflow Locks"),
        ("#covenant-compliance", "Covenant Compliance"),
        ("#forensic-audit-dashboard", "Forensic Audit – Overview"),
        ("#math-integrity", "Forensic Audit – Math Integrity"),
        ("#performance-benchmarking", "Forensic Audit – Performance"),
        ("#fraud-detection", "Forensic Audit – Fraud Detection"),
        ("#tenant-risk", "Forensic Audit – Tenant Risk"),
        ("#collections-quality", "Forensic Audit – Collections"),
        ("#document-completeness", "Forensic Audit – Documents"),
        ("#reconciliation-results", "Forensic Audit – Reconciliation"),
        ("#audit-history", "Forensic Audit – History"),
        ("#nlq-search", "Natural Language Query"),
        ("#rule-configuration/{id}", "Rule Configuration"),
        ("#rule-edit/{id}", "Rule Edit"),
        ("#syntax-guide", "Syntax Guide"),
    ]
    for route, page in route_list:
        r = routes_table.add_row()
        r.cells[0].text = route
        r.cells[1].text = page
    doc.add_paragraph()

    # What to look at - Financial Integrity Hub
    doc.add_heading("What End Users Should Look At – Financial Integrity Hub", level=1)
    doc.add_paragraph("When using the Financial Integrity Hub (#forensic-reconciliation), focus on:")
    for item in [
        "Integrity Score – Ensure it is in the Pass (green) range before closing",
        "Overall Status – Resolve any Fail (red) before month-end",
        "Critical Issues (Overview tab) – Address BS-1, IS-1, and large variances first",
        "Exceptions tab – Triage and resolve all Critical/High items",
        "Covenant Compliance – Verify DSCR, LTV within thresholds",
        "By Rule / By Document – Drill into specific failures",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()

    # Maintenance
    doc.add_heading("Document Information", level=1)
    doc.add_paragraph("Last Updated: January 31, 2026")
    doc.add_paragraph("Version: 1.0")
    doc.add_paragraph("Replace all screenshot placeholders with actual captures from your REIMS instance for best end-user experience.")

    out_dir = Path("/home/hsthind/REIMS - Documents")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "REIMS2_Complete_System_User_Manual.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return str(out_path)


if __name__ == "__main__":
    main()
