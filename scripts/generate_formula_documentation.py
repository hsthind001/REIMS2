#!/usr/bin/env python3
"""
Generate REIMS2 Financial Formulas Documentation

Creates a .docx documenting how each key financial metric is calculated,
which line items and document types are used, and formula details.

Output: /home/hsthind/REIMS - Documents/REIMS2_Financial_Formulas_Reference.docx
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

FORMULAS = [
    {
        "name": "Property Value",
        "alternate_names": ["Current Value", "Net Property Value"],
        "formula": "Gross Property Value - Accumulated Depreciation\nOR Total Assets (fallback)",
        "calculation": (
            "Gross Property Value = Land (0510-0000) + Buildings (0610-0000) + Improvements (0710-0950)\n"
            "Accumulated Depreciation = Sum of 1061-1091 (contra-asset, negative)\n"
            "Net Property Value = Gross Property + Accumulated Depreciation (accum_depr is negative)\n\n"
            "Fallback (LTV/Cap Rate services): Total Assets = Sum of all Balance Sheet accounts starting with '1' (1xxxx)\n"
            "Last resort: total_area_sqft × $200/sf or $5M default"
        ),
        "line_items": [
            "0510-0000 Land",
            "0610-0000 Buildings",
            "0710-0950 Improvements (5 Year, 15 Year, Roof, HVAC, TI, etc.)",
            "1061-1091 Accumulated Depreciation (Buildings, 5 Year, 15 Year, Roof, Other)",
            "1999-0000 Total Assets (fallback)",
        ],
        "document_types": ["Balance Sheet"],
        "source_tables": ["balance_sheet_data", "financial_metrics"],
    },
    {
        "name": "Property NOI",
        "alternate_names": ["NOI", "Net Operating Income"],
        "formula": "NOI = Total Revenue - Operating Expenses",
        "calculation": (
            "Total Revenue = Sum of Income Statement accounts 4xxxx (Total Income/Revenue)\n"
            "Operating Expenses = Sum of 5xxxx and 6xxxx accounts ONLY\n"
            "Excludes: 7xxxx (Mortgage Interest), 8xxxx (Depreciation, Amortization)\n\n"
            "If monthly data: Annualized NOI = NOI × 12 (for DSCR/cap rate)\n"
            "Prefer calculated NOI; fallback to stored NET OPERATING INCOME line (6299-0000)"
        ),
        "line_items": [
            "4999-0000 Total Income / Total Revenue",
            "5xxxx Operating expenses (Property Tax, Insurance, Management, etc.)",
            "6xxxx Additional operating expenses",
            "6299-0000 NET OPERATING INCOME (if stored)",
        ],
        "document_types": ["Income Statement"],
        "source_tables": ["income_statement_data", "financial_metrics"],
    },
    {
        "name": "Occupancy Rate",
        "formula": "Occupancy Rate = (Occupied Units / Total Units) × 100",
        "calculation": (
            "Total Units = Count of leasable units (excludes COMMON, ATM, LAND, SIGN)\n"
            "Occupied Units = Count where occupancy_status = 'occupied'\n"
            "Alternative (by sqft): Occupied Sqft / Total Leasable Sqft × 100"
        ),
        "line_items": [
            "unit_number (to filter leasable)",
            "occupancy_status (occupied, vacant, notice)",
            "unit_area_sqft",
            "monthly_rent",
        ],
        "document_types": ["Rent Roll"],
        "source_tables": ["rent_roll_data", "financial_metrics"],
    },
    {
        "name": "Property DSCR",
        "alternate_names": ["DSCR", "Debt Service Coverage Ratio"],
        "formula": "DSCR = Annualized NOI / Annual Debt Service",
        "calculation": (
            "Annualized NOI = NOI × (12 / window_months) if monthly data\n"
            "Annual Debt Service = Monthly (Principal + Interest) × 12\n\n"
            "Sources for debt service:\n"
            "1. FinancialMetrics.total_annual_debt_service (preferred)\n"
            "2. Mortgage Statement: principal_due + interest_due, or total_payment_due\n"
            "3. Income Statement interest (7010) + estimated principal\n\n"
            "Covenant thresholds: ≥1.50 Strong, 1.25-1.49 Adequate, 1.00-1.14 Critical, <1.00 Default Risk"
        ),
        "line_items": [
            "Income Statement: NET OPERATING INCOME, TOTAL INCOME, TOTAL EXPENSE",
            "Mortgage: total_payment_due, principal_due, interest_due, tax_escrow_due, insurance_escrow_due, reserve_due",
        ],
        "document_types": ["Income Statement", "Mortgage Statement"],
        "source_tables": ["income_statement_data", "mortgage_statement_data", "financial_metrics"],
    },
    {
        "name": "LTV",
        "alternate_names": ["Loan-to-Value", "LTV Ratio"],
        "formula": "LTV = Current Loan Balance / Net Property Value",
        "calculation": (
            "Loan Balance (preferred): Sum of principal_balance from mortgage_statement_data\n"
            "Loan Balance (fallback): Balance Sheet 2900-0000 Total Long-Term Liabilities\n\n"
            "Net Property Value: 1099-0000 Total Property & Equipment, or Gross (0510-0950) - Accum Depr (1061-1091)\n\n"
            "Thresholds: ≤65% Strong, 65-75% Good, 75-80% Fair, >80% High Risk"
        ),
        "line_items": [
            "Mortgage: principal_balance (preferred)",
            "Balance Sheet: 2900-0000 Total Long-Term Liabilities (fallback)",
            "Property: 1099-0000 Total Property & Equipment, or net_property_value (Gross - Accum Depr)",
        ],
        "document_types": ["Mortgage Statement", "Balance Sheet"],
        "source_tables": ["mortgage_statement_data", "balance_sheet_data", "financial_metrics"],
    },
    {
        "name": "Purchase Price",
        "formula": "Stored value from Property record",
        "calculation": (
            "Purchase Price = Original acquisition price stored in properties.purchase_price\n"
            "Total Acquisition Cost = purchase_price + acquisition_costs (closing costs, legal fees, due diligence)"
        ),
        "line_items": ["N/A – stored in properties table"],
        "document_types": ["Property Record (non-document)"],
        "source_tables": ["properties"],
    },
    {
        "name": "Current Value",
        "formula": "Same as Property Value",
        "calculation": (
            "Current Value = Most recent period's Total Assets or Net Property Value\n"
            "Used interchangeably with Property Value in cap rate, LTV, exit strategy calculations"
        ),
        "line_items": ["See Property Value"],
        "document_types": ["Balance Sheet"],
        "source_tables": ["balance_sheet_data", "financial_metrics"],
    },
    {
        "name": "Hold Period",
        "formula": "Hold Period = Current Date - Acquisition Date",
        "calculation": (
            "Investment Date = properties.acquisition_date\n"
            "Hold Period (years) = (Current Date - Acquisition Date) in days / 365.25\n"
            "Used in exit strategy analysis, IRR calculations"
        ),
        "line_items": ["acquisition_date from properties table"],
        "document_types": ["Property Record"],
        "source_tables": ["properties"],
    },
    {
        "name": "Cap Rate",
        "alternate_names": ["Capitalization Rate"],
        "formula": "Cap Rate = NOI / Property Value\nProperty Value (estimate) = NOI / Cap Rate",
        "calculation": (
            "Cap Rate = Net Operating Income / Property Value (decimal, e.g., 0.075 = 7.5%)\n\n"
            "For valuation: Property Value = Annualized NOI / Target Cap Rate\n"
            "If monthly: Annualized NOI = NOI × 12 (or ×4 for quarterly)\n\n"
            "Market benchmarks by type: Retail 7.5%, Office 7%, Mixed-Use 8%, Industrial 6.5%, Multifamily 6%"
        ),
        "line_items": [
            "NOI from Income Statement (4xxxx revenue - 5xxxx/6xxxx expenses)",
            "Property Value from Balance Sheet (total_assets or net_property_value)",
        ],
        "document_types": ["Income Statement", "Balance Sheet"],
        "source_tables": ["income_statement_data", "balance_sheet_data", "financial_metrics"],
    },
]


def add_formula_section(doc, f):
    """Add a formula section to the document."""
    names = f["name"]
    if f.get("alternate_names"):
        names += f" ({', '.join(f['alternate_names'])})"
    doc.add_heading(names, level=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    def row(label, value):
        r = table.add_row()
        r.cells[0].text = label
        r.cells[1].text = str(value) if value else "—"
        r.cells[0].paragraphs[0].runs[0].bold = True

    row("Formula", f["formula"])
    row("Calculation Details", f["calculation"])
    row("Line Items / Account Codes", "\n".join(f["line_items"]))
    row("Document Types", ", ".join(f["document_types"]))
    row("Source Tables", ", ".join(f["source_tables"]))
    doc.add_paragraph()


def main():
    doc = Document()
    doc.add_heading("REIMS2 Financial Formulas Reference", 0)
    doc.add_paragraph()
    doc.add_paragraph(
        "This document provides detailed information on how key financial metrics are calculated in REIMS2, "
        "including formulas, line items from source documents, and document types used."
    )
    doc.add_paragraph()

    doc.add_heading("Table of Contents", level=1)
    for f in FORMULAS:
        doc.add_paragraph(f"• {f['name']}", style="List Bullet")
    doc.add_paragraph()

    doc.add_heading("Formula Details", level=1)
    for f in FORMULAS:
        add_formula_section(doc, f)

    doc.add_heading("Document Type Reference", level=1)
    ref_table = doc.add_table(rows=0, cols=2)
    ref_table.style = "Table Grid"
    h = ref_table.add_row()
    h.cells[0].text = "Document Type"
    h.cells[1].text = "Key Account Ranges / Fields"
    for c in h.cells:
        c.paragraphs[0].runs[0].bold = True
    refs = [
        ("Balance Sheet", "1xxxx Assets; 2xxxx Liabilities; 3xxxx Equity. Totals: 1999, 2999, 3999. Property: 0510 Land, 0610 Buildings, 0710-0950 Improvements, 1061-1091 Accum Depr."),
        ("Income Statement", "4xxxx Revenue; 5xxxx-6xxxx Operating Expenses; 7xxxx Interest; 8xxxx Depr/Amort; 6299 NOI; 9090 Net Income."),
        ("Mortgage Statement", "principal_balance, principal_due, interest_due, total_payment_due, tax_escrow_due, insurance_escrow_due."),
        ("Rent Roll", "unit_number, occupancy_status, unit_area_sqft, monthly_rent, annual_rent."),
    ]
    for dt, codes in refs:
        r = ref_table.add_row()
        r.cells[0].text = dt
        r.cells[1].text = codes
    doc.add_paragraph()

    doc.add_heading("Maintenance", level=1)
    doc.add_paragraph("Last Updated: January 31, 2026")
    doc.add_paragraph("Document Version: 1.0")

    out_dir = Path("/home/hsthind/REIMS - Documents")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "REIMS2_Financial_Formulas_Reference.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return str(out_path)


if __name__ == "__main__":
    main()
