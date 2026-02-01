#!/usr/bin/env python3
"""
Generate Financial Integrity Hub User Manual

Creates a comprehensive .docx user manual for the Financial Integrity Hub,
including purpose, components, tabs, and end-user workflows.

Output: /home/hsthind/REIMS - Documents/REIMS2_Financial_Integrity_Hub_User_Manual.docx
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO_ROOT = Path(__file__).resolve().parent.parent
# Use REAL screenshots from this folder. See SCREENSHOT_CAPTURE_GUIDE.md for capture instructions.
SCREENSHOTS_DIR = Path("/home/hsthind/REIMS - Documents/reims-screenshots")


def add_screenshot(doc, filename, caption, width=5.5):
    """Add real screenshot if it exists; otherwise add placeholder with capture instructions."""
    path = SCREENSHOTS_DIR / filename
    if path.exists():
        try:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(path), width=Inches(width))
            if caption:
                cap = doc.add_paragraph()
                cap.add_run(caption).italic = True
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            doc.add_paragraph(f"[Insert screenshot: {filename}]")
    else:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("⚠️ INSERT REAL SCREENSHOT: ").bold = True
        p.add_run(f"Capture from the live REIMS app. Save as '{filename}' in reims-screenshots folder. See SCREENSHOT_CAPTURE_GUIDE.md for exact steps.")
        p.paragraph_format.space_before = Pt(6)
    doc.add_paragraph()


def add_heading_with_decision(doc, title, content, decisions=None):
    """Add a section with optional 'Decisions you can make' callout."""
    doc.add_heading(title, level=2)
    doc.add_paragraph(content)
    if decisions:
        p = doc.add_paragraph()
        p.add_run("Decisions you can make: ").bold = True
        p.add_run(decisions)
        p.paragraph_format.space_before = Pt(6)
    doc.add_paragraph()


def main():
    doc = Document()
    doc.add_heading("REIMS2 Financial Integrity Hub", 0)
    doc.add_paragraph()
    doc.add_heading("User Manual", level=1)
    doc.add_paragraph(
        "This manual explains how to use the Financial Integrity Hub to validate financial data, "
        "run reconciliation rules, and make informed decisions about your property's financial health."
    )
    doc.add_paragraph()

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "Purpose of the Financial Integrity Hub",
        "How to Access the Financial Integrity Hub",
        "Hub Layout Overview",
        "Header & Controls (Property, Period, Run Reconciliation, Validate, Export)",
        "Metrics Grid (Integrity Score, Overall Status, Reconciliation Stats)",
        "Reconciliation Matrix",
        "Tabs: Overview, By Document, By Rule, Exceptions, Insights",
        "Modals & Panels (Document Pair Panel, Edit Rule Modal)",
        "What End Users Should Look At",
        "Recommended Workflow",
    ]
    for item in toc_items:
        doc.add_paragraph(f"• {item}", style="List Bullet")
    doc.add_paragraph()

    # Purpose
    doc.add_heading("Purpose of the Financial Integrity Hub", level=1)
    doc.add_paragraph(
        "The Financial Integrity Hub is REIMS2's central dashboard for validating and reconciling "
        "financial data across all document types. It enables you to:"
    )
    bullets = [
        "Validate that Balance Sheet, Income Statement, Cash Flow, Rent Roll, and Mortgage Statement data are internally consistent",
        "Run 300+ automated reconciliation rules to detect variances and anomalies",
        "Monitor covenant compliance (DSCR, LTV, liquidity, occupancy)",
        "Identify critical issues (e.g., Accounting Equation failures, large variances) before month-end close",
        "Track cross-document matches and discrepancies between related line items",
        "Export validation results for audit and compliance reporting",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()

    # How to Access
    doc.add_heading("How to Access the Financial Integrity Hub", level=1)
    doc.add_paragraph(
        "Navigate to Operations → Financial Integrity Hub (or use the route #forensic-reconciliation). "
        "You can also access it from the Financials page or Quality Control page via the \"Open Financial Integrity Hub\" button."
    )
    doc.add_paragraph()

    # Layout Visual
    doc.add_heading("Hub Layout Overview", level=1)
    add_screenshot(doc, "financial-integrity-hub-main.png", "Figure 1: Financial Integrity Hub main page")
    doc.add_paragraph()

    # Header & Controls
    doc.add_heading("Header & Controls", level=1)
    doc.add_paragraph()
    add_heading_with_decision(
        doc,
        "Property Selector",
        "Select the property you want to validate. The hub shows reconciliation results for the selected property only. Your selection is saved for your next visit.",
        "Which property needs validation this period?",
    )
    add_heading_with_decision(
        doc,
        "Period Selector",
        "Select the financial period (e.g., 2025-01 for January 2025). Periods are loaded based on the selected property. The most recent period is often auto-selected.",
        "Which reporting period should I validate?",
    )
    add_heading_with_decision(
        doc,
        "Run Reconciliation",
        "Click \"Run Reconciliation\" to execute all reconciliation rules for the selected property and period. This creates a session, runs matches, and calculates rule results. Wait 10–30 seconds for completion.",
        "When to run: After uploading new documents, before month-end close, or when refreshing validation.",
    )
    add_heading_with_decision(
        doc,
        "Validate",
        "Click \"Validate\" to recalculate the health score from existing matches. Run this after \"Run Reconciliation\" to update the Integrity Score.",
        "When to run: After reconciliation completes to refresh the overall health metric.",
    )
    add_heading_with_decision(
        doc,
        "Export",
        "Export validation results for audit trails or external review.",
        "When to use: For compliance documentation or sharing with auditors.",
    )
    doc.add_paragraph()

    # Metrics Grid
    doc.add_heading("Metrics Grid (Top Cards)", level=1)
    doc.add_paragraph()
    add_heading_with_decision(
        doc,
        "Integrity Score",
        "A 0–100% gauge showing overall financial data health. Scores ≥90% = Pass, 70–89% = Warning, <70% = Fail. Based on rule pass rates and match quality.",
        "Is my data healthy enough to close the books? Do I need to investigate further?",
    )
    add_heading_with_decision(
        doc,
        "Overall Status",
        "Traffic light: Pass (green), Warning (yellow), or Fail (red). Also shows the count of active rules (e.g., 377 Rules Active).",
        "Quick yes/no: Can I proceed or must I fix issues first?",
    )
    add_heading_with_decision(
        doc,
        "Reconciliation Stats",
        "Verified Matches: Count of cross-document matches that passed validation. Discrepancies: Count of items that need review. Click Discrepancies to jump to the Exceptions tab.",
        "How many items are reconciled vs. need attention?",
    )
    doc.add_paragraph()

    # Reconciliation Matrix
    doc.add_heading("Reconciliation Matrix", level=1)
    doc.add_paragraph(
        "A 5×5 grid showing cross-document reconciliation results. Rows and columns represent: Balance Sheet, Income Statement, Cash Flow, Rent Roll, Mortgage Statement. "
        "Each cell shows the number of matches between that document pair (e.g., Balance Sheet ↔ Income Statement). "
        "Cells are color-coded: green for passed, amber for variances, gray for no data. "
        "Click a cell to open the Document Pair Panel with detailed match lists."
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Decisions you can make: ").bold = True
    p.add_run("Which document pairs have the most variances? Where should I focus my review?")
    doc.add_paragraph()

    # Tabs
    doc.add_heading("Tabs & Tab Content", level=1)
    doc.add_paragraph()

    tabs = [
        {
            "name": "Overview",
            "icon": "📊",
            "content": (
                "High-level dashboard with: (1) Document Health – health score per document type (Balance Sheet, Income Statement, etc.) with progress bars; "
                "(2) Covenant Compliance – DSCR, LTV, liquidity, occupancy status; (3) Critical Issues – top discrepancies and rule violations with amounts; "
                "(4) Recent Activity – match/alert timestamps; (5) Covenant History – compliance by period; (6) Variance Alerts – AUDIT-48 variance breach alerts."
            ),
            "decisions": "Which documents are weakest? Are covenants in compliance? What are the top issues to fix?",
        },
        {
            "name": "By Document",
            "icon": "📄",
            "content": (
                "Rules grouped by document type. Expand each document (Balance Sheet, Income Statement, Cash Flow, Rent Roll, Mortgage Statement, Three-Statement Integration) "
                "to see its rules with Pass/Fail status, formula, actual value, and severity. Helps you understand which document type has the most failures."
            ),
            "decisions": "Which document type has the most rule failures? Should I re-extract or correct data for a specific document?",
        },
        {
            "name": "By Rule",
            "icon": "📋",
            "content": (
                "All rules in a filterable list. Summary cards: Total Rules, Passed, Variance, Skipped, Pass Rate. Filter by status (click cards). Search by rule ID or name. "
                "Each rule card shows status, rule name, formula, and values. Click a rule to open the Edit Logic modal (if authorized) or view configuration."
            ),
            "decisions": "Which specific rules are failing? What is my pass rate? Should I adjust a rule threshold or fix the source data?",
        },
        {
            "name": "Exceptions",
            "icon": "🚨",
            "content": (
                "Unified list of all exceptions: forensic discrepancies and rule violations. Sorted by severity (Critical → High → Medium → Low). "
                "Shows description, source, amount difference, and suggested resolution. Filter and triage issues before month-end close."
            ),
            "decisions": "What are all the issues I need to resolve? In what order should I address them? Are there any critical blockers?",
        },
        {
            "name": "Insights",
            "icon": "💡",
            "content": (
                "Analytical insights: Reconciliation Efficiency (pass rate vs. baseline), Best/Worst performing document types, and dynamic recommendations. "
                "Summarizes matches and discrepancies with actionable suggestions (e.g., \"Balance Sheet needs attention – 3 rules failed\")."
            ),
            "decisions": "What are the trends? Which areas need improvement? What actions are recommended?",
        },
    ]

    for t in tabs:
        add_heading_with_decision(doc, f"{t['icon']} {t['name']} Tab", t["content"], t["decisions"])

    # Modals & Panels
    doc.add_heading("Modals & Panels", level=1)
    doc.add_paragraph()
    doc.add_heading("Document Pair Panel", level=2)
    doc.add_paragraph(
        "Opens when you click a cell in the Reconciliation Matrix. Shows reconciliation details for that document pair (e.g., Balance Sheet → Income Statement): "
        "passed checks count, variances count, and a list of individual matches. Use it to drill into specific cross-document validations."
    )
    doc.add_paragraph()
    doc.add_heading("Edit Rule Modal", level=2)
    doc.add_paragraph(
        "Opens when you click a rule in the By Rule tab (or via Configure Rule). Allows authorized users to edit the rule's formula, threshold, or description. "
        "Changes apply to future reconciliation runs."
    )
    doc.add_paragraph()

    # What to Look At
    doc.add_heading("What End Users Should Look At", level=1)
    doc.add_paragraph(
        "Before month-end close or audit, focus on:"
    )
    checklist = [
        "Integrity Score – ensure it is in the Pass (green) range",
        "Overall Status – resolve any Fail (red) before closing",
        "Critical Issues (Overview tab) – address Accounting Equation (BS-1), Net Income (IS-1), and large variances first",
        "Exceptions tab – triage by severity and resolve or document all Critical/High items",
        "Covenant Compliance – verify DSCR, LTV, and other covenants are within thresholds",
        "Variance Alerts – review AUDIT-48 breach alerts and investigate material variances",
    ]
    for c in checklist:
        doc.add_paragraph(c, style="List Bullet")
    doc.add_paragraph()

    # Workflow
    doc.add_heading("Recommended Workflow", level=1)
    doc.add_paragraph("1. Select Property and Period")
    doc.add_paragraph("2. Click Run Reconciliation and wait for completion")
    doc.add_paragraph("3. Click Validate to refresh the Integrity Score")
    doc.add_paragraph("4. Review Overview tab for Document Health and Critical Issues")
    doc.add_paragraph("5. Open Exceptions tab and resolve Critical/High items")
    doc.add_paragraph("6. Check Covenant Compliance and Variance Alerts")
    doc.add_paragraph("7. Use By Rule or By Document tabs to drill into specific failures")
    doc.add_paragraph("8. Export results if needed for audit")
    doc.add_paragraph()

    # Maintenance
    doc.add_heading("Maintenance", level=1)
    doc.add_paragraph("Last Updated: January 31, 2026")
    doc.add_paragraph("Document Version: 1.0")

    out_dir = Path("/home/hsthind/REIMS - Documents")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "REIMS2_Financial_Integrity_Hub_User_Manual.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return str(out_path)


if __name__ == "__main__":
    main()
